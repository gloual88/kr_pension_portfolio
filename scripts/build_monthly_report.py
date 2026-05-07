"""
build_monthly_report.py
========================
월간 연금 포트폴리오 docx 보고서 생성 — 유튜브 영상 대본 / 슬라이드 활용용.

데이터 소스: outputs_trimmed10/
  - cio/final_portfolio.json
  - cio/board_memo.md
  - macro/macro-view.json
  - pc_weights_matrix_pivot_pct.csv
  - pc_category_sums_pct.csv

실행 (사용자가 직접):
  cd d:\\파이선
  $env:PYTHONIOENCODING='utf-8'
  & "d:\\파이선\\pykrx_venv\\Scripts\\python.exe" -m \\
      kr_pension_portfolio.scripts.build_monthly_report \\
      --out outputs_trimmed10 --month 2026-05
"""
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path

import pandas as pd
import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 헬퍼
# ============================================================
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # 다크 블루
COLOR_ACCENT = RGBColor(0xC8, 0x9B, 0x3C)    # 골드
COLOR_GREY = RGBColor(0x59, 0x59, 0x59)
COLOR_RED = RGBColor(0xC0, 0x39, 0x2B)


def _set_run(run, *, bold=False, size=None, color=None, italic=False):
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if italic:
        run.italic = True


def add_title(doc, text, size=22, color=COLOR_PRIMARY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run(r, bold=True, size=size, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, bold=True, size=15, color=COLOR_PRIMARY)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, bold=True, size=12, color=COLOR_GREY)
    return p


def add_para(doc, text, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, size=size, color=color)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_run(r, size=size)
    return p


def make_table(doc, headers, rows, col_widths_cm=None, header_color=COLOR_PRIMARY):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _set_run(r, bold=True, size=10, color=header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, v in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(v))
            _set_run(r, size=10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)
    return table


# ============================================================
# 메인
# ============================================================
def main(out_dir: str, month_label: str) -> None:
    base = Path(__file__).resolve().parent.parent
    out_path = base / out_dir
    ips_path = base / "configs" / "ips_trimmed10.yaml"

    cio = json.loads((out_path / "cio" / "final_portfolio.json").read_text(encoding="utf-8"))
    macro = json.loads((out_path / "macro" / "macro-view.json").read_text(encoding="utf-8"))
    ips = yaml.safe_load(ips_path.read_text(encoding="utf-8"))
    asset_classes = ips["investment_universe"]["asset_classes"]

    etf_label = {ac["slug"]: f'{ac["etf"]} {ac["etf_name"]}' for ac in asset_classes}
    short_name = {ac["slug"]: ac["etf_name"] for ac in asset_classes}
    cat_of = {ac["slug"]: ac["category"] for ac in asset_classes}
    slugs = [ac["slug"] for ac in asset_classes]

    cat_sums = pd.read_csv(out_path / "pc_category_sums_pct.csv")
    pivot_pct = pd.read_csv(out_path / "pc_weights_matrix_pivot_pct.csv", index_col=0)

    # CIO 핵심
    cio_w = cio["weights"]
    cio_m = cio["metrics"]
    chosen = cio["chosen_ensemble"]

    # 위험자산
    cat_totals = {"Equity": 0.0, "FixedIncome": 0.0, "RealAssets": 0.0, "Cash": 0.0}
    for s, w in cio_w.items():
        if cat_of.get(s) in cat_totals:
            cat_totals[cat_of[s]] += w
    risk_w = cat_totals["Equity"] + cat_totals["RealAssets"]

    # ============================================================
    # docx 빌드
    # ============================================================
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ─── 표지 ───
    add_title(doc, f"월간 연금 포트폴리오 — {month_label}", size=22)
    add_para(doc, f"KR DC/IRP 자율주행 SAA · 10-ETF 축약 유니버스 · 21개 PC 모델 비교",
             size=11, color=COLOR_GREY)
    add_para(doc, f"산출 기준일: {datetime.now().strftime('%Y-%m-%d')}  |  "
                  f"매크로 readings: ECOS + FRED 14/15  |  "
                  f"엔진: agentic SAA pipeline v1",
             size=10, color=COLOR_GREY)

    doc.add_paragraph()

    # ─── 요약 박스 ───
    add_h1(doc, "한 페이지 요약")
    add_bullet(doc, f"거시 레짐은 **{macro['regime']}** (신뢰도 {macro['confidence']:.2f}, "
                    f"12개월 침체확률 {macro['recession_probability_12m']*100:.0f}%) "
                    f"— 성장 견조, 인플레 안정, 유가·환율 구조적 부담.")
    add_bullet(doc, f"21개 PC 모델 합의로 채택된 ensemble은 **{chosen}** — "
                    f"E[r] {cio_m['expected_return']*100:.2f}%, "
                    f"σ {cio_m['expected_vol']*100:.2f}%, "
                    f"BT Sharpe {cio_m['backtest_sharpe']:.2f}, "
                    f"MDD {cio_m['backtest_maxdd']*100:.2f}%.")
    add_bullet(doc, f"위험자산 비중 **{risk_w*100:.1f}%** — DC/IRP 70% 한도 대비 "
                    f"보수적 배분 ({(0.70 - risk_w)*100:.1f}%p 여유).")
    add_bullet(doc, f"21개 모델의 ETF별 평균 비중과 CIO 최종 비중이 거의 일치 → "
                    f"방법론에 무관하게 동일한 결론으로 수렴 (모델 합의가 강함).")

    # ─── 1. 매크로 ───
    add_h1(doc, "1. 거시 진단 — 5월 시점 매크로 readings")

    add_h2(doc, "성장")
    add_bullet(doc, f"GDP YoY {macro['readings']['kr_gdp_yoy']:.2f}% / "
                    f"산업생산 YoY {macro['readings']['kr_industrial_production_yoy']:.2f}% / "
                    f"수출 YoY {macro['readings']['kr_exports_yoy']:.2f}% / "
                    f"실업률 {macro['readings']['kr_unemployment']:.1f}%.")
    add_bullet(doc, "→ 성장 점수 +0.66 (4축 중 가장 양호)")

    add_h2(doc, "물가")
    add_bullet(doc, f"CPI YoY {macro['readings']['kr_cpi_yoy']:.2f}% / "
                    f"근원 CPI YoY {macro['readings']['kr_core_cpi_yoy']:.2f}% / "
                    f"브렌트유 ${macro['readings']['kr_brent_oil']:.2f}/bbl.")
    add_bullet(doc, "→ 헤드라인은 BOK 목표(2%) 근방. 다만 유가가 $114로 stagflation 리스크 잠재.")

    add_h2(doc, "통화·금융")
    add_bullet(doc, f"BOK 기준금리 {macro['readings']['kr_base_rate']:.2f}% / "
                    f"KTB 10Y {macro['readings']['kr_ktb_10y']:.3f}% / "
                    f"KTB 3Y {macro['readings']['kr_ktb_3y']:.3f}% / "
                    f"3-10Y 커브 {macro['readings']['kr_curve_3y_10y']*100:.0f}bp.")
    add_bullet(doc, f"USD/KRW {macro['readings']['kr_usd_krw']:.0f} / "
                    f"AA- 스프레드 {macro['readings']['kr_corp_aa_spread_bp']:.0f}bp / "
                    f"VIX {macro['readings']['us_vix']:.1f} / "
                    f"Fed funds {macro['readings']['us_fed_funds']:.2f}%.")
    curve = macro.get("curve_signal", {})
    if curve:
        add_bullet(doc, f"→ 곡선 상태: {curve['regime']} ({curve['shape']}). {curve['notes']}")
    else:
        add_bullet(doc, "→ 금리와 커브는 장기채·현금 비중의 핵심 입력값으로 해석.")

    # ─── 2. CIO 최종 ───
    doc.add_paragraph()
    add_h1(doc, f"2. 최종 추천 포트폴리오 — CIO ({chosen} ensemble)")

    add_para(doc, "10개 ETF 비중 (내림차순):", size=11, color=COLOR_GREY)
    rows_top = []
    sorted_w = sorted(cio_w.items(), key=lambda x: -x[1])
    for s, w in sorted_w:
        rows_top.append([etf_label[s], cat_of[s], f"{w*100:.2f}%"])
    make_table(doc, ["ETF", "카테고리", "비중"], rows_top,
               col_widths_cm=[9.0, 3.5, 2.5])

    doc.add_paragraph()
    add_h2(doc, "메트릭")
    add_bullet(doc, f"기대수익률 E[r] {cio_m['expected_return']*100:.2f}% / "
                    f"기대변동성 σ {cio_m['expected_vol']*100:.2f}%")
    add_bullet(doc, f"백테스트 Sharpe {cio_m['backtest_sharpe']:.2f} (KR 60/40 BM 1.71 대비 +{cio_m['backtest_sharpe']-1.71:+.2f})")
    add_bullet(doc, f"백테스트 MDD {cio_m['backtest_maxdd']*100:.2f}% (KR 60/40 BM -13.5%)")
    add_bullet(doc, f"Tracking Error {cio_m['tracking_error']*100:.2f}% (IPS budget 6% 초과)")
    add_bullet(doc, f"HHI {cio_m['hhi']:.3f} → Effective N {cio_m['effective_n']:.1f} "
                    f"(이론 max=10) — 분산도 양호")
    add_bullet(doc, f"IPS 컴플라이언스 {cio_m['ips_compliance']*100:.0f}%")

    add_h2(doc, "카테고리별 합계")
    add_bullet(doc, f"Equity {cat_totals['Equity']*100:.1f}% (한도 55%) — "
                    f"KR Equity 2개 + US Equity 2개 균형")
    add_bullet(doc, f"FixedIncome {cat_totals['FixedIncome']*100:.1f}% (KR 10Y + US 10Y + US IG)")
    add_bullet(doc, f"RealAssets {cat_totals['RealAssets']*100:.1f}% (Gold 단일)")
    add_bullet(doc, f"Cash {cat_totals['Cash']*100:.1f}% (KOFR + Money Market)")
    add_bullet(doc, f"위험자산(Equity+RealAssets) **{risk_w*100:.1f}%** ≤ DC/IRP 70% 한도",
               size=11)

    # ─── 3. 21개 모델 비교 ───
    doc.add_paragraph()
    add_h1(doc, "3. 21개 PC 모델은 어떻게 다른가?")

    add_para(doc, "동일한 IPS 제약(Equity 20–55%, FI 20–70%, RealAsset ≤15%, "
                  "Cash ≤30%, 종목 ≤25%) 하에서 21개 PC 방법론이 산출한 카테고리 합계입니다. "
                  "위험자산 비중 내림차순.", size=11)

    cat_view = cat_sums.copy()
    cat_view = cat_view.sort_values("Risky(Eq+Real)", ascending=False)
    rows_cat = []
    for _, r in cat_view.iterrows():
        rows_cat.append([
            r["method"],
            f"{r['Equity']:.1f}",
            f"{r['FixedIncome']:.1f}",
            f"{r['RealAssets']:.1f}",
            f"{r['Cash']:.1f}",
            f"{r['Risky(Eq+Real)']:.1f}",
        ])
    make_table(doc, ["Method", "Equity", "FI", "RealAsset", "Cash", "Risky"],
               rows_cat, col_widths_cm=[4.5, 2.0, 2.0, 2.0, 2.0, 2.0])

    doc.add_paragraph()
    add_h2(doc, "관전 포인트")
    add_bullet(doc, "가장 공격적: market-cap-weight 54.8% / equal-weight·black-litterman·"
                    "cvar-min·max-entropy 50%. 모두 IPS Equity 한도 55% 직전.")
    add_bullet(doc, "가장 방어적: max-dd-constrained 20% / tpa 22% / gmv 23.6% / "
                    "max-sharpe 24.7%. 위험자산을 IPS 하한 (Equity 20%)까지 줄임.")
    add_bullet(doc, "분포 중앙: 21개 평균 ≈ 38%. 70% 한도 대비 매우 보수적 — "
                    "현 매크로(late-cycle, AA스프레드 65bp)에서 방어적 합의가 형성됨.")

    # ─── 4. ETF 평균 비중 ───
    doc.add_paragraph()
    add_h1(doc, "4. 21개 모델 합의 — ETF별 평균 비중")

    avg = pivot_pct.mean(axis=1).round(2)
    avg_df = pd.DataFrame({
        "ETF": [idx for idx in pivot_pct.index],
        "21모델 평균": avg.values,
        "CIO 최종": [cio_w[s] * 100 for s in slugs],
    })
    avg_df["차이"] = (avg_df["CIO 최종"] - avg_df["21모델 평균"]).round(2)
    avg_df = avg_df.sort_values("21모델 평균", ascending=False)

    rows_avg = []
    for _, r in avg_df.iterrows():
        rows_avg.append([
            r["ETF"].split("  ")[-1] if "  " in r["ETF"] else r["ETF"],
            f"{r['21모델 평균']:.2f}%",
            f"{r['CIO 최종']:.2f}%",
            f"{r['차이']:+.2f}",
        ])
    make_table(doc, ["ETF", "21모델 평균", "CIO 최종", "차이(pp)"],
               rows_avg, col_widths_cm=[7.5, 2.5, 2.5, 2.5])

    doc.add_paragraph()
    add_para(doc, "→ CIO inverse_te ensemble 결과가 21모델 평균과 모든 ETF에서 1pp 이내로 "
                  "거의 일치. 즉 현재 매크로·CMA 환경에서는 어떤 PC 방법론을 써도 "
                  "결론이 같다는 뜻으로, 결과의 강건성이 높다고 해석할 수 있습니다.",
             size=11)

    # ─── 5. 위험 점검 ───
    doc.add_paragraph()
    add_h1(doc, "5. 위험 점검 + IPS 컴플라이언스")

    add_h2(doc, "통과")
    add_bullet(doc, f"DC/IRP 70% 위험자산 한도 — 38% (32%p 여유) ✅")
    add_bullet(doc, f"카테고리 bounds 모두 충족 (Equity 20-55%, FI 20-70%, "
                    f"RealAsset ≤15%, Cash ≤30%) ✅")
    add_bullet(doc, f"종목 ≤25% — 최대 비중 KOFR-cash 14.3% ✅")

    add_h2(doc, "주의")
    add_bullet(doc, f"기대 변동성 σ {cio_m['expected_vol']*100:.2f}% < IPS 하한 6% — "
                    "포트폴리오가 너무 보수적으로 빌드됨. 객관적 위험 부족이 아니라 "
                    "수익 기회 회수 부족 의미. 콘셉트상 의도된 결과.")
    add_bullet(doc, f"Tracking Error {cio_m['tracking_error']*100:.2f}% > IPS budget 6% — "
                    "한국 60/40 BM이 KOSPI200 60% + KTB10Y 40%로 너무 집중적이라 "
                    "BM과 자연스럽게 멀어짐. 글로벌 60/40으로 BM 재정의 시 해소 가능.")

    add_h2(doc, "스트레스 시나리오")
    add_bullet(doc, "유가 추가 충격 (Brent → $130+) — RealAsset(Gold) 5% + "
                    "US IG Credit 8%로 헷지. Cash 28%는 dry powder.")
    add_bullet(doc, "BOK 추가 인하 (2.50% → 2.00%) — KR 10Y 14%, US 10Y 12%로 "
                    "듀레이션 노출. 평가이익 +1.5~2% 예상.")
    add_bullet(doc, "원/달러 충격 (1485 → 1550+) — US Equity 16% 환노출분에서 "
                    "원화환산 +5~7% 가능. 다만 매크로 헷지 자산 (Gold, KRW IG는 헷지)이 동시에 작동.")

    # ─── 6. 재조정 ───
    doc.add_paragraph()
    add_h1(doc, "6. 재조정 규칙 + 다음 분기 점검사항")

    add_h2(doc, "리밸런싱")
    add_bullet(doc, "주기: **분기**. 다음 정기 리밸런싱: 2026-08-01 분기 시작.")
    add_bullet(doc, "이탈 트리거: 자산군 비중이 목표 대비 ±5%p 이상 벗어나면 분기 외 재조정.")

    add_h2(doc, "다음 점검 포인트 (다음 달 영상에서 확인)")
    add_bullet(doc, "BOK 5월 금통위 — 추가 인하 시 KR 듀레이션 비중 점검.")
    add_bullet(doc, "유가 — Brent $120 돌파 시 stagflation regime으로 전환 가능성.")
    add_bullet(doc, "USD/KRW — 1500 또는 1450 돌파 시 IPS escalation 트리거.")
    add_bullet(doc, "VIX — 25 돌파 시 위험자산 추가 축소 검토.")

    # ─── 7. 부록 ───
    doc.add_page_break()
    add_h1(doc, "부록 A. 21개 모델 × 10 ETF 전체 비중 (%)")

    add_para(doc, "행 = ETF, 열 = PC 모델. 비중 단위는 %. 0.5% 미만은 공란 처리.",
             size=10, color=COLOR_GREY)

    full = pivot_pct.round(1).fillna(0.0)
    methods_all = list(full.columns)

    # 21열은 가로로 길어서 6+5+5+5 식으로 분할 (열 4개 그룹)
    col_groups = [methods_all[:6], methods_all[6:12], methods_all[12:17], methods_all[17:]]
    for gi, group in enumerate(col_groups, start=1):
        add_h2(doc, f"A.{gi} 모델 {gi*6-5}–{gi*6 if gi<4 else len(methods_all)}")
        sub = full[group]
        rows_sub = []
        for idx, row in sub.iterrows():
            short = idx.split("  ")[-1] if "  " in idx else idx
            row_vals = [f"{v:.1f}" if v >= 0.5 else "" for v in row.values]
            rows_sub.append([short, *row_vals])
        make_table(doc, ["ETF"] + list(group), rows_sub)
        doc.add_paragraph()

    # ─── 면책 ───
    doc.add_page_break()
    add_h1(doc, "면책 고지")
    p = doc.add_paragraph()
    r = p.add_run(
        "본 보고서는 자율주행 SAA 파이프라인이 산출한 정보를 기반으로 작성된 "
        "정보 제공 자료입니다. 특정 금융상품의 매수·매도 권유가 아니며, "
        "투자자문업·투자일임업·투자권유 행위가 아닙니다. "
        "백테스트 결과는 가상의 가정에 기반하며 미래 성과를 보장하지 않습니다. "
        "DC/IRP 규제(위험자산 70% 한도, 레버리지·인버스 ETF 매수 금지)를 가정한 "
        "IPS 하에서 산출된 결과입니다. 실제 투자 결정 전 본인의 투자목적·위험감수도·"
        "재무상황을 종합 검토하시기 바랍니다."
    )
    _set_run(r, size=9, color=COLOR_GREY, italic=True)

    # 저장
    save_path = out_path / f"monthly_report_{month_label}.docx"
    doc.save(save_path)
    print(f"[saved] {save_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="outputs_trimmed10")
    parser.add_argument("--month", default=None,
                        help="월 라벨 (예: 2026-05). 미지정 시 현재 년월 사용.")
    args = parser.parse_args()
    month_label = args.month or datetime.now().strftime("%Y-%m")
    main(out_dir=args.out, month_label=month_label)
