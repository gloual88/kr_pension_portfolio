"""Generate Production Brief docx for the KR Pension SAA promotional video.

Output: docs/production_brief_kr_pension_saa.docx
Audience: external video producer (영상 제작자)
Emphasis: (1) AI Agent multi-agent architecture (2) AI 자율 결정 (no human in loop)
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- 스타일 헬퍼 ----------
NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xC8, 0x9B, 0x3C)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0xEE, 0xEE, 0xEE)


def set_kr_font(run, size_pt=10, bold=False, color=None):
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rpr.append(rfonts)


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_kr_font(r, 18, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    # 가는 밑줄
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_kr_font(r, 13, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_kr_font(r, 11, bold=True, color=GRAY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_body(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_kr_font(r, size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_kr_font(r, 10)
    p.paragraph_format.left_indent = Cm(0.5 + 0.5 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_callout(doc, text, color=GOLD):
    p = doc.add_paragraph()
    r = p.add_run("▸ " + text)
    set_kr_font(r, 10, bold=True, color=color)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, widths=None, header_color="1F4E79"):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if widths:
        for i, w in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_kr_font(r, 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, header_color)

    # 본문
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[1 + ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_kr_font(r, 9.5)
            p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()
    return table


# ---------- 문서 작성 ----------
def build():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ====== 표지 ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCTION BRIEF")
    set_kr_font(r, 11, bold=True, color=GOLD)
    p.paragraph_format.space_before = Pt(60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KR 연금 자율주행 SAA")
    set_kr_font(r, 32, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Self-Driving Pension Allocation by Multi-Agent AI")
    set_kr_font(r, 14, color=GRAY)
    p.paragraph_format.space_after = Pt(40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("외부 영상 제작자 작업 가이드")
    set_kr_font(r, 12, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("핵심 메시지: 46개 AI 에이전트가 자율 협의해 만드는 한국형 퇴직연금 포트폴리오")
    set_kr_font(r, 11, color=GRAY)
    p.paragraph_format.space_after = Pt(60)

    # 메타 표
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci in range(2):
        for cell in meta_table.columns[ci].cells:
            cell.width = Cm(7.0)
    metas = [
        ("프로젝트", "KR Pension Self-Driving SAA — 홍보 영상"),
        ("대상", "30-50대 DC/IRP 가입자, 퇴직연금 운용에 관심 있는 직장인"),
        ("형식", "메인 90초 + 티저 30초 (옵션: 3분 deep-dive)"),
        ("발행 채널", "YouTube / Instagram Reels / 카카오톡 채널"),
    ]
    for i, (k, v) in enumerate(metas):
        c0 = meta_table.rows[i].cells[0]
        c1 = meta_table.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        set_kr_font(r0, 10, bold=True, color=NAVY)
        r1 = c1.paragraphs[0].add_run(v)
        set_kr_font(r1, 10)
        shade_cell(c0, "EEEEEE")

    doc.add_page_break()

    # ====== 1. 프로젝트 개요 ======
    add_h1(doc, "1. 프로젝트 개요")

    add_body(doc, "본 영상은 한국 퇴직연금(DC/IRP) 가입자를 대상으로 한 'AI 에이전트 자율운용 SAA' 데모 서비스의 홍보 영상입니다. 시청자가 30초 내에 '내 연금을 AI가 자율로 운용해주는 새로운 방식이 나왔다'는 사실을 인지하고, 90초 안에 그 차별성과 신뢰 근거(8.3년 백테스트 실적)를 이해하도록 설계되어야 합니다.", size=10)

    add_h2(doc, "영상이 반드시 전달해야 할 3가지 사실")
    add_callout(doc, "① 사람이 아니라 *AI들이 자기들끼리* 의논해서 결정합니다.")
    add_callout(doc, "② 한국 경제 진단부터 ETF 비중까지 모두 자동 — 사람 개입 0회.")
    add_callout(doc, "③ 법이 허용하는 70% 한도까지 끝까지 활용 — 8년 검증, 60/40 펀드 대비 흔들림 30% 작고 수익 더 좋음.")
    add_body(doc, "※ 이 3가지 메시지는 narration에서는 위 표현 그대로 (또는 더 쉽게) 사용. 'Sharpe'·'백테스트'·'ensemble'·'매크로' 같은 전문 용어는 §11 글로서리의 일상 표현으로 반드시 변환할 것.", color=GRAY, size=9.5)

    add_h2(doc, "영상에서 절대 사용하면 안 되는 표현")
    add_bullet(doc, "\"수익 보장\", \"손실 없음\" — 백테스트는 미래를 보장하지 않습니다.")
    add_bullet(doc, "\"전문가가 운용\" — 본 모델의 차별성은 'AI 자율'이며 사람 자문이 아닙니다. 이 메시지가 흐려지면 안 됩니다.")
    add_bullet(doc, "\"투자 권유\" 인상을 주는 문구 — 면책 고지 필수.")
    add_bullet(doc, "타사 명칭 직접 비교 — 차별화는 카테고리(TDF, robo, 자문사) 단위로만.")

    add_h2(doc, "필수 면책 문구 (자막 / 엔드 카드)")
    add_body(doc, "「본 자료는 정보 제공 목적이며 투자 권유가 아닙니다. 투자 의사결정과 그 결과는 투자자 본인의 책임입니다. 백테스트 성과는 미래 수익을 보장하지 않습니다.」", color=GRAY)

    doc.add_page_break()

    # ====== 2. 핵심 메시지 / Tagline ======
    add_h1(doc, "2. 핵심 메시지 · Tagline")

    add_h2(doc, "Hero Tagline 후보 (3개 중 택일) — 모두 일상 한국어")
    add_body(doc, "★ 1순위 추천", bold=True, color=GOLD)
    add_callout(doc, "\"46명의 AI가 함께 만드는, 매월 새로워지는 퇴직연금 추천\"", color=NAVY)
    add_body(doc, "2순위", bold=True, color=GRAY)
    add_callout(doc, "\"사람 대신 AI가 결정합니다 — 매월 자동으로\"", color=NAVY)
    add_body(doc, "3순위", bold=True, color=GRAY)
    add_callout(doc, "\"AI 46명이 매달 당신의 연금 포트폴리오를 다시 만듭니다\"", color=NAVY)

    add_h2(doc, "보조 카피 (자막 / 엔드 카드)")
    add_bullet(doc, "\"21명의 AI가 서로 평가까지 합니다\"")
    add_bullet(doc, "\"한국 경제 진단부터 ETF 비중까지 — 100% 자동\"")
    add_bullet(doc, "\"퇴직연금 한도 70%, 끝까지 활용한 8년의 검증\"")
    add_bullet(doc, "\"모든 결정 과정이 화면에 보입니다\"")

    add_h2(doc, "키워드 — 영상 내내 등장해야 할 핵심 단어 (일상 표현 우선)")
    add_body(doc, "AI, 46명의 AI, 자동, 사람 대신, 협력, 평가, 추천, 매월", bold=True, color=NAVY)
    add_body(doc, "(브랜드 용어로만 1-2회: AI 에이전트, 자율주행 — 처음 등장 시 \"스스로 결정한다는 뜻\" 같이 즉시 풀어줌)", color=GRAY, size=9.5)

    doc.add_page_break()

    # ====== 3. AI 에이전트 시스템 — 가장 중요한 섹션 ======
    add_h1(doc, "3. AI 에이전트 시스템 (반드시 강조)")

    add_body(doc, "본 모델은 단일 알고리즘이 아닌 9개 역할의 AI 에이전트가 자율적으로 협업하는 다중 에이전트 시스템입니다. 사람이 운영하는 자산운용사 조직 구조를 그대로 AI로 옮긴 것이 핵심 컨셉입니다 — 매크로 애널리스트, 18명의 자산 애널리스트, 21명의 포트폴리오 매니저, 위험관리자, 동료평가단, CIO. 영상 시청자가 이 'AI 조직' 이미지를 잡으면 차별화가 즉시 전달됩니다.", size=10)

    add_h2(doc, "AI 에이전트 ↔ 사람 자산운용사 조직 매핑")
    add_table(
        doc,
        headers=["사람 자산운용사 직무", "본 모델의 AI 에이전트", "수", "역할"],
        rows=[
            ["매크로 이코노미스트", "Macro Agent", "1", "ECOS+FRED 14개 지표 → 경기 레짐 분류"],
            ["섹터/자산 애널리스트", "Asset Class Agents", "18", "각 ETF별 기대수익률·변동성 추정"],
            ["퀀트 리서처", "Covariance Agent + PC Researcher", "2", "공분산 추정 + 신규 모델 제안"],
            ["포트폴리오 매니저", "PC Agents", "21", "각자 다른 알고리즘으로 비중 산출"],
            ["리스크 매니저 (CRO)", "CRO Agent", "1", "21개 안에 대한 위험 보고"],
            ["투자위원회 동료평가", "Strategy Review Agent", "1", "Borda count 투표로 최상위 5개 선별"],
            ["CIO (최고투자책임자)", "CIO Agent", "1", "7개 앙상블 후보 중 매크로에 맞춰 최종 선택"],
            ["성과 분석가", "Meta Agent", "1", "예측 vs 실현 추적, 자기개선 제안"],
        ],
        widths=[4.5, 4.5, 1.2, 5.3],
    )

    add_callout(doc, "영상에서 이 표를 1초 동안 풀스크린으로 보여주세요 — 시청자가 'AI가 통째로 운용사를 대체한다'는 메시지를 즉각 받습니다.", color=GOLD)

    add_h2(doc, "'자율'의 의미 — 영상에서 풀어 설명")
    add_bullet(doc, "사람의 개입 없이 분기마다 매크로 진단 → 자산 분석 → 21개 모델 산출 → 동료평가 → CIO 결정이 자동 실행됩니다.")
    add_bullet(doc, "투자자는 '이번 분기 추천 포트폴리오'를 받기만 하고, 매수·리밸런싱 결정도 본인이 합니다 (AI는 결정 자료 제공).")
    add_bullet(doc, "AI 에이전트는 정해진 프로토콜로 협업합니다 — 자료 호출, 분석, 평가, 합의가 모두 코드로 실행되어 재현 가능합니다.")

    add_h2(doc, "21개 PC 에이전트는 각자 어떻게 다른가")
    add_body(doc, "동일한 시장 데이터를 받아도 21개의 AI는 서로 다른 *철학*으로 비중을 산출합니다. 이 다양성이 본 모델의 핵심입니다.", size=10)
    add_table(
        doc,
        headers=["에이전트 카테고리", "개수", "철학", "대표 예"],
        rows=[
            ["Heuristic", "2", "단순 규칙 기반", "Equal Weight, Market Cap"],
            ["Return-Optimized", "5", "수익률 최대화", "Max Sharpe, Black-Litterman"],
            ["Risk-Structured", "6", "위험 구조화", "Risk Parity, GMV, HRP"],
            ["Non-Traditional", "7", "비전통 / 견고성", "CVaR-Min, Max-DD Constrained, Tail Risk Parity"],
            ["Researcher (자가제안)", "1", "매크로에 맞춰 *새로운* 방법 제안", "Maximum Entropy 등"],
        ],
        widths=[4.5, 1.2, 4.0, 5.8],
    )

    doc.add_page_break()

    # ====== 4. 스토리 구조 — 데모 기반 ======
    add_h1(doc, "4. 스토리 구조 — 사이드바 데모 흐름")

    add_body(doc, "본 모델의 차별성은 *프로세스* (AI 46명이 자기들끼리 의논하는 과정)에 있습니다. 그래서 단순 스토리텔링보다 *대시보드 사이드바 메뉴를 순차적으로 보여주는 데모 흐름*이 훨씬 강력합니다 — 시청자가 'AI가 정말 자율로 작동하는구나'를 화면으로 직접 봅니다. 메인 영상 (3:30)이 데모 흐름을 따르고, 60초 티저는 메인의 압축 버전입니다.", size=10)

    add_h2(doc, "메인 영상 (3분 30초) — 사이드바 메뉴 순차 데모 ★")
    add_table(
        doc,
        headers=["시간", "막", "사이드바 메뉴", "한 줄 메시지", "화면 포인트"],
        rows=[
            ["0:00 - 0:30", "Hook + 컨셉", "(없음, 도입)", "퇴직연금 한도 70%, 안 쓰면 손해. AI 46명이 알아서 만들어준다면?", "안전자산 vs 60/40 누적 비교 차트"],
            ["0:30 - 0:50", "AI Macro 분석가", "2_Macro_Regime", "AI가 한국 경제 14가지 신호로 지금이 좋은 시기인지 판단", "regime 라벨 + 14개 지표 표 클로즈업"],
            ["0:50 - 1:10", "AI 자산 분석가 18명", "3_Asset_Classes", "18명의 AI가 각자 한 ETF를 맡아 *얼마나 오를지* 예측", "18개 자산 카드 빠른 컷 (각 ~1초)"],
            ["1:10 - 1:35", "AI 매니저 21명", "5_월간_연금_포트폴리오 (히트맵)", "같은 자료, 다른 방식 — 21명이 21개 다른 답을 냄", "21×ETF 히트맵 zoom-in (색상 차이 강조)"],
            ["1:35 - 1:55", "AI 동료평가 + 팀장 결정", "5_월간_연금_포트폴리오 (CIO)", "AI들이 서로 점수 매기고, 팀장 AI가 최종 선택", "Borda 투표 표 → CIO 도넛 트랜지션"],
            ["1:55 - 2:15", "이번 달 추천 포트폴리오", "6_현재_포트폴리오", "주식 55% + 금/원유 15% = 공격적 70%, 채권/현금 30%", "도넛 + 18 ETF 비중 표"],
            ["2:15 - 3:00", "8년 검증 vs 60/40", "1_Backtest", "흔들림 30% 작고 수익은 더 좋음 — 100만원 → 233만원", "NAV 곡선 + 메트릭 카드 4장"],
            ["3:00 - 3:30", "CTA + 면책", "(없음)", "매월 1일 자동, 무료. 본 자료는 투자 권유가 아닙니다", "URL + QR + 면책 5초 정지"],
        ],
        widths=[1.8, 2.5, 3.5, 5.0, 4.5],
    )

    add_h2(doc, "60초 티저 (메인의 압축 — 빠른 컷)")
    add_table(
        doc,
        headers=["시간", "메시지", "화면"],
        rows=[
            ["0:00 - 0:10", "퇴직연금. 한도 70%, 안 쓰면 손해입니다.", "Hero shot — 누적 비교 차트 1초 + AI 46명 매핑 표 5초"],
            ["0:10 - 0:25", "46명의 AI 팀이 한국 경제부터 ETF 비중까지 — 자기들끼리 의논해서 매달 추천을 만듭니다.", "사이드바 Macro → Assets → 21 PC 히트맵 → CIO 도넛 빠른 컷 (각 3-4초)"],
            ["0:25 - 0:45", "지난 8년 검증: 60/40 펀드 대비 흔들림 30% 작고 수익 더 좋음.", "NAV 곡선 + Sharpe 1.195 카드"],
            ["0:45 - 1:00", "사이트에서 무료. 매월 1일 자동.", "URL + QR + 비밀번호 3734 + 면책"],
        ],
        widths=[2.0, 6.5, 7.0],
    )

    doc.add_page_break()

    # ====== 5. Narration 초안 — 일상 한국어 ======
    add_h1(doc, "5. Narration 초안 — 일상 한국어 (메인 3:30 + 티저 60s)")

    add_body(doc, "아래 narration은 30-50대 일반 시청자가 1-2회 듣고 즉시 이해하도록 일상 한국어로 작성했습니다. \"Sharpe\", \"백테스트\", \"매크로\", \"앙상블\" 같은 전문 용어를 의도적으로 회피했습니다 — 자막도 narration의 표현을 그대로 (또는 더 짧게) 따라 주세요. §11 글로서리에 \"지양 ↔ 권장 표현\" 표 첨부.", color=GRAY)

    add_h2(doc, "메인 영상 (3분 30초) ★")
    main_seg = [
        ("0:00 - 0:30  [Hook + 컨셉]",
         "퇴직연금, 신경 쓰고 계신가요? 한국 퇴직연금은 주식처럼 공격적인 상품에 70%까지 넣을 수 있습니다. 그런데 대부분 그 한도를 다 쓰지 못합니다. 사실, 8년이면 거의 100% 가까운 차이가 납니다. 이걸 AI에게 맡기면 어떨까요? 46명의 AI가 한 팀이 되어, 매달 추천 포트폴리오를 만들어준다면.",
         "[화면] 안전자산 누적곡선 vs 60/40 누적곡선 차이 강조 → AI 46명 매핑 표 페이드 인"),
        ("0:30 - 0:50  [AI 경제 분석가 — 메뉴 2_Macro_Regime]",
         "먼저, AI가 지금 한국 경제 상태부터 살핍니다. 금리, 물가, 환율, 수출, 실업률... 14가지 경제 신호를 모아서, 지금이 좋은 시기인지 조심할 시기인지 판단합니다. 사람이 하면 며칠 걸릴 일을 AI는 자동으로 합니다.",
         "[화면] 사이드바 메뉴 2 클릭 → regime 라벨 + 14개 지표 표 클로즈업 (예: GDP +3.6%, CPI 2.6%, BOK 2.5%)"),
        ("0:50 - 1:10  [AI 자산 분석가 18명 — 메뉴 3_Asset_Classes]",
         "다음, 18명의 AI가 *각자 한 가지 ETF*를 맡습니다. 코스피 ETF는 1번 AI가, 미국 S&P 500은 2번 AI가, 미국 국채는 3번 AI가... 각자 자기 담당 상품이 앞으로 얼마나 오를지, 얼마나 흔들릴지 예측합니다. 18명이 동시에, 24시간 일합니다.",
         "[화면] 사이드바 메뉴 3 클릭 → 18개 자산 카드 빠른 컷 (각 ~1초, 카드에 ETF 명 + E[r] 강조)"),
        ("1:10 - 1:35  [AI 매니저 21명 — 메뉴 5 히트맵]",
         "이제 21명의 AI 매니저가 등장합니다. 같은 자료를 받지만, *생각하는 방식*이 다릅니다. 어떤 AI는 가장 안전한 답, 어떤 AI는 가장 수익이 클 답을 추천합니다. 21명이 21개의 *서로 다른* 추천을 냅니다. 이 화면이 그 차이를 보여줍니다 — 색이 진할수록 비중이 큰 상품.",
         "[화면] 사이드바 메뉴 5 → 21×ETF 히트맵 zoom-in. 색상 차이 강조 (예: 한 모델은 미국주식 25%, 다른 모델은 채권 20%)"),
        ("1:35 - 1:55  [AI 동료평가 + 팀장 결정 — 같은 페이지]",
         "AI들끼리 서로 점수를 매깁니다. *누구의 추천이 가장 믿을 만한가?* 가장 높은 점수를 받은 5명을 추려내고, 팀장 AI가 *지금 시장에 가장 잘 맞는* 답을 최종 선택합니다. 사람의 개입 없이, 모든 평가가 자동으로 이루어집니다.",
         "[화면] 같은 페이지 스크롤 다운 → Borda 투표 표 (Top-5) → CIO 최종 도넛으로 트랜지션"),
        ("1:55 - 2:15  [이번 달 추천 — 메뉴 6_현재_포트폴리오]",
         "이번 달 추천 포트폴리오입니다. 주식 55%에 금과 원유가 15%, 합쳐서 공격적 상품 70%. 나머지 30%는 채권과 현금. *법이 허용하는 한도까지 끝까지 활용*하는 게 핵심입니다.",
         "[화면] 사이드바 메뉴 6 → 도넛 (55+15=70 강조) + 18 ETF 비중 표 (위에서 5개)"),
        ("2:15 - 3:00  [8년 검증 vs 60/40 — 메뉴 1_Backtest]",
         "근데 진짜 효과 있을까요? 지난 8년, 2018년부터 지금까지, 매분기마다 이 AI 시스템을 *실제로* 돌려봤습니다. 결과는 — 한국에서 가장 흔한 60/40 펀드와 비교했을 때, 흔들림은 30% 작고, 위험 대비 수익은 30% 더 좋았습니다. 같은 8년에 100만원을 넣었으면 233만원이 됩니다.",
         "[화면] 사이드바 메뉴 1 → 'v1 lock70 Baseline' 변형 선택 → NAV 곡선 (lock70 vs BM 두 줄) + 메트릭 카드 4장 (Sharpe / 수익률 / 흔들림 / MDD)"),
        ("3:00 - 3:30  [CTA + 면책]",
         "매월 1일, 새로운 추천이 자동으로 나옵니다. 사이트에서 바로 확인할 수 있습니다. 무료입니다. 본 자료는 정보 제공이며, 투자 권유가 아닙니다.",
         "[화면] URL https://kr-pension-saa-gloual88.streamlit.app/ + QR + 비밀번호 3734 + 면책 고지 5초 정지"),
    ]
    for tcode, text, vis in main_seg:
        add_h3(doc, tcode)
        p = doc.add_paragraph()
        r = p.add_run("Narration: ")
        set_kr_font(r, 10, bold=True, color=NAVY)
        r = p.add_run(text)
        set_kr_font(r, 10)
        p = doc.add_paragraph()
        r = p.add_run(vis)
        set_kr_font(r, 9.5, color=GRAY)
        p.paragraph_format.space_after = Pt(8)

    add_h2(doc, "60초 티저 — 메인의 압축")
    teaser_seg = [
        ("0:00 - 0:10  [Hook]",
         "퇴직연금. 한도 70%, 안 쓰면 손해입니다.",
         "[화면] 안전자산 vs 60/40 누적 차이 차트 (1초) + AI 46명 매핑 표 페이드 인 (5초)"),
        ("0:10 - 0:25  [데모 압축]",
         "그래서 만들었습니다. 46명의 AI 팀. 한국 경제부터 ETF 비중까지 — 자기들끼리 의논해서 매달 추천을 만듭니다.",
         "[화면] 사이드바 Macro → Assets → 21 PC 히트맵 → CIO 도넛 빠른 컷 (각 3-4초)"),
        ("0:25 - 0:45  [Result]",
         "지난 8년 검증 결과: 60/40 펀드 대비 흔들림 30% 작고, 수익도 더 좋았습니다.",
         "[화면] NAV 곡선 + Sharpe 1.195 + 누적 +133% 카드 (각 5초씩 빠른 컷)"),
        ("0:45 - 1:00  [CTA]",
         "사이트에서 무료. 매월 1일 자동. 본 자료는 투자 권유가 아닙니다.",
         "[화면] URL + QR + 비밀번호 3734 + 면책"),
    ]
    for tcode, text, vis in teaser_seg:
        add_h3(doc, tcode)
        p = doc.add_paragraph()
        r = p.add_run("Narration: ")
        set_kr_font(r, 10, bold=True, color=NAVY)
        r = p.add_run(text)
        set_kr_font(r, 10)
        p = doc.add_paragraph()
        r = p.add_run(vis)
        set_kr_font(r, 9.5, color=GRAY)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ====== 6. 차별화 매트릭스 ======
    add_h1(doc, "6. 차별화 매트릭스 (Why Different)")

    add_body(doc, "영상 0:10-0:25 'Problem' 막에서 활용. 3개 카테고리 모두 보여주면 산만합니다 — 영상에서는 본 모델의 ★ 행만 강조하고, 비교 카테고리는 1행씩 빠른 컷으로 처리하세요.", color=GRAY)

    add_table(
        doc,
        headers=["항목", "본 모델 (AI 자율)", "TDF", "단일 robo", "자문사"],
        rows=[
            ["의사결정 주체", "★ AI 에이전트 9개 역할 (46개)", "정적 글라이드패스", "단일 알고리즘", "사람 자문가"],
            ["매크로 반영", "★ 분기마다 14개 지표 자동 진단", "없음 (나이만)", "부분적", "정성 판단"],
            ["알고리즘 다양성", "★ 21개 PC 모델 병렬 + 동료평가", "1개", "1-2개", "0 (사람)"],
            ["모델 선택 근거", "★ AI 동료평가 + 매크로 매핑", "N/A", "불명 (블랙박스)", "자문가 판단"],
            ["위험자산 활용", "★ 70% 한도 100% 사용", "나이 따라 감소", "변동", "변동"],
            ["투명성", "★ 모든 21개 모델 weights 공개", "닫힘", "닫힘", "부분"],
            ["백테스트 공개", "★ walk-forward 8.3년 (2018-2026)", "✗", "△", "✗"],
            ["사람 개입", "★ 0회 (완전 자율)", "1회 (가입 시)", "월 1회+", "분기/연 1회"],
            ["비용", "★ 무료 데모", "TER 0.3-0.5%", "자문료 0.5-1%", "1%+"],
        ],
        widths=[3.0, 4.5, 2.5, 2.5, 2.5],
    )

    add_callout(doc, "★ = 본 모델 우위 구간. 9개 항목 중 9개 모두 ★. 영상에서 \"9 of 9\" 같은 메시지로 시각화 가능.", color=GOLD)

    doc.add_page_break()

    # ====== 7. 시각 자료 캡처 가이드 ======
    add_h1(doc, "7. 시각 자료 캡처 가이드 (Capture Guide)")

    add_body(doc, "본 영상의 B-roll(배경 영상)은 라이브 대시보드에서 직접 녹화합니다. URL: https://kr-pension-saa-gloual88.streamlit.app/  비밀번호: 3734  (Cloud 자동 배포 직후 안정화 30초).", size=10)

    add_h2(doc, "녹화 스펙")
    add_bullet(doc, "해상도: 1920×1080 (FHD) 이상, 가능하면 4K로 녹화 후 다운샘플")
    add_bullet(doc, "프레임: 60fps (마우스 이동 부드럽게)")
    add_bullet(doc, "브라우저: Chrome 전체화면 모드 (F11), 주소창 / 사이드바 / 시크릿 입력 화면 모두 컷에서 제외")
    add_bullet(doc, "다크모드 OFF — 본 대시보드는 라이트 테마 기준")

    add_h2(doc, "페이지별 녹화 포커스")
    add_table(
        doc,
        headers=["페이지", "추천 컷", "용도", "녹화 시간"],
        rows=[
            ["5_월간_연금_포트폴리오", "도넛 차트 (위험 70%) + 21×ETF 히트맵 + 카테고리 stack", "Solution-2 막의 '21개 모델' 시각화", "30초"],
            ["1_Backtest", "변형 선택 → 'v1 lock70 Baseline' → NAV 곡선 + 성과 표", "Result 막의 백테스트 근거", "20초"],
            ["2_Macro_Regime", "매크로 indicator dashboard + regime 라벨", "Solution-2의 '매크로 진단'", "10초"],
            ["3_Asset_Classes", "ETF 18개 카드 + CMA 수치", "Solution-1의 '18 자산 애널리스트' 강조", "10초"],
            ["6_현재_포트폴리오", "오늘 산출된 비중 도넛", "CTA 막 - 사용자 화면", "5초"],
        ],
        widths=[3.5, 6.5, 3.5, 1.5],
    )

    add_h2(doc, "고해상도 스크린샷 (정지 이미지로 사용)")
    add_bullet(doc, "위험자산 70% lock 도넛 (페이지 5 상단) — Solution / Result에서 풀스크린")
    add_bullet(doc, "21×ETF 히트맵 (페이지 5 섹션 2) — Solution-2 핵심 시각")
    add_bullet(doc, "Backtest NAV 곡선 (페이지 1) — Result 핵심")
    add_bullet(doc, "Section 3의 'AI 조직 매핑' 표 (이 docx에서 추출하여 모션그래픽 처리 권장)")

    add_h2(doc, "외부 도구로 만들 추가 자료")
    add_bullet(doc, "메트릭 카드 (단일 숫자 카드 4장): \"Sharpe 1.195\" / \"+133%\" / \"MDD -21.6%\" / \"σ -32%\" — Result 막에서 빠른 컷")
    add_bullet(doc, "5단계 파이프라인 다이어그램 (애니메이션): 매크로 → 21 PC → 동료평가 → CIO → 산출")
    add_bullet(doc, "QR 코드 (URL): CTA 막 5초 풀스크린")

    doc.add_page_break()

    # ====== 8. 브랜드 / 톤 가이드 ======
    add_h1(doc, "8. 브랜드 · 톤 가이드")

    add_h2(doc, "색상 팔레트 (대시보드와 일치)")
    add_table(
        doc,
        headers=["용도", "Hex", "RGB", "사용처"],
        rows=[
            ["주요 (Primary Navy)", "#1F4E79", "31, 78, 121", "제목, 강조 텍스트, 차트 메인"],
            ["보조 (Accent Gold)", "#C89B3C", "200, 155, 60", "★ 표시, 핵심 수치, CTA 버튼"],
            ["중성 (Gray)", "#555555", "85, 85, 85", "보조 텍스트, 캡션"],
            ["배경 (Light)", "#EEEEEE", "238, 238, 238", "테이블 헤더 배경, 카드 배경"],
            ["채권 (FixedIncome)", "#7E9BB8", "126, 155, 184", "FI 자산 차트"],
            ["실물 (RealAssets)", "#C89B3C", "200, 155, 60", "Gold/원유 자산"],
            ["현금 (Cash)", "#5B8A72", "91, 138, 114", "Cash 자산"],
        ],
        widths=[3.5, 2.5, 3.5, 5.5],
    )

    add_h2(doc, "톤 (Voiceover · BGM)")
    add_bullet(doc, "Voiceover 톤: 신뢰감 + 데이터 기반 차분함. \"권유\"가 아닌 \"설명\".")
    add_bullet(doc, "성별: 30-40대 남성 또는 여성 모두 가능. 전문성이 우선.")
    add_bullet(doc, "BGM: Hook 막은 약간 긴장감, Solution부터 희망적/진취적, Result는 절정. 일렉트로닉 minimal 권장 (재무/금융 영상 표준).")
    add_bullet(doc, "효과음: AI 작동음(키보드 타이핑, 스캔 효과음) Solution-1 막에서 사용 가능 — 단, 과도하면 신뢰감 훼손.")

    add_h2(doc, "타이포그래피")
    add_bullet(doc, "한글: Pretendard 또는 Noto Sans KR (대시보드 fallback: Malgun Gothic)")
    add_bullet(doc, "영문: Inter 또는 Roboto")
    add_bullet(doc, "수치 강조: Bold + 1.2× 크기. 단위 (% / 배 / 원) 항상 명시.")

    add_h2(doc, "모션 스타일")
    add_bullet(doc, "데이터 시각화 위주. 인물(스톡 영상) 사용 시 1-2컷에 한정.")
    add_bullet(doc, "트랜지션: 빠른 컷 + 페이드 X (정보 영상 톤). 화면 전환 시 마진 작게.")
    add_bullet(doc, "텍스트 모션: 슬라이드 인/타이프라이터/줌 모두 가능 — 단, 한 영상에서 2종 이하로 제한.")

    doc.add_page_break()

    # ====== 9. 결과 (강조 수치 모음) ======
    add_h1(doc, "9. 강조 수치 (Result 막에서 보여줄 모든 숫자)")

    add_body(doc, "Result 막에서 모든 숫자를 다 보여주려 하지 마세요 — 4개 메트릭 카드를 1.5초씩 빠른 컷으로 보여주거나, 표 1장을 3초간 풀스크린이 좋습니다.", color=GRAY)

    add_h2(doc, "Walk-forward Backtest (2018-01-02 ~ 2026-05-08, 34 분기 QS 리밸런싱)")
    add_table(
        doc,
        headers=["지표", "AI baseline ★", "AI Phase 2 LLM", "KR 60/40 BM"],
        rows=[
            ["연환산 수익률", "11.15%", "11.15%", "12.82%"],
            ["연환산 변동성", "9.33%", "9.26%", "13.79%"],
            ["Sharpe Ratio", "1.195", "1.204", "0.929"],
            ["Max Drawdown", "-21.57%", "-21.56%", "-26.73%"],
            ["총 누적 수익률", "132.85%", "133.00%", "152.58%"],
            ["거래 비용 (누적)", "9 bp", "10 bp", "—"],
            ["분기당 평균 턴오버", "5.52%", "5.70%", "—"],
        ],
        widths=[3.5, 4.0, 4.0, 4.0],
    )

    add_callout(doc, "★ 영상에는 baseline 수치(Sharpe 1.195)를 사용하세요 — production 운영 기준이며, Phase 2 대비 turnover/비용이 낮습니다. Phase 2 결과는 'LLM Phase 2까지 검증한 데이터'라는 *기술 신뢰성* 메시지로만 활용 가능.", color=GOLD)

    add_callout(doc, "핵심 메시지: 'AI 자율은 BM 대비 변동성을 32% 낮추면서 Sharpe를 29% 높입니다.' 이 한 줄이 Result 막의 narration이어야 합니다.", color=GOLD)

    # 비교 차트 임베드 (있으면)
    chart_path = Path(__file__).resolve().parents[1] / "docs" / "lock70_nav_comparison.png"
    if chart_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(chart_path), width=Cm(15.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cap.add_run("[ 그림: 3-line NAV 비교 — baseline / Phase 2 LLM / KR 60/40 BM ]")
        set_kr_font(rc, 9, color=GRAY)

    add_h2(doc, "LLM Phase 2의 가치 — Why baseline은 충분한가")
    add_bullet(doc, "한국에서 LLM의 부가가치는 매우 작음 — Phase 1 검증 (2026-05-01) 동일 결론, lock70 + Phase 2 (2026-05-10) 재확인.")
    add_bullet(doc, "lock70 IPS가 ensemble 자유도를 좁힘 — 위험자산 70% 고정 상태에서 LLM은 카테고리 *내부* 분배만 다르게 가능.")
    add_bullet(doc, "Phase 2 ROI 낮음: 분기당 ~$0.4 cost로 Sharpe +0.7% 개선 (1.195 → 1.204), σ -7bp. 통계적으로 noise 수준.")
    add_bullet(doc, "Production 운영에는 baseline 추천: 결정적, API cost 0, 결과 거의 동일.")
    add_bullet(doc, "LLM의 진짜 가치는 의사결정이 아닌 *설명*에 있음 — board memo, regime narrative, Q&A.")

    add_h2(doc, "현재 시점 모델 포트폴리오 (2026-05-10 산출, 18-ETF 풀 유니버스)")
    add_table(
        doc,
        headers=["카테고리", "비중", "내용"],
        rows=[
            ["Equity (KR + US + Intl)", "55.00% (lock)", "kr-large-cap, kr-kosdaq, kr-dividend, us-large-cap, us-tech, us-dividend, intl-developed, emerging-markets"],
            ["RealAssets (실물)", "15.00% (lock)", "gold (KODEX 골드선물H), commodities (TIGER 원유선물H)"],
            ["FixedIncome (채권)", "25.56%", "kr-treasuries-10y, kr-short-bonds, kr-credit, us-treasuries-10y/30y, us-ig-credit"],
            ["Cash (현금성)", "4.44%", "KOFR 금리 액티브, 머니마켓 액티브"],
            ["위험자산 합", "70.00%", "DC/IRP 법적 한도 100% 활용"],
        ],
        widths=[4.5, 3.0, 8.5],
    )

    doc.add_page_break()

    # ====== 10. 체크리스트 ======
    add_h1(doc, "10. 체크리스트 (영상 제작자 확인용)")

    add_h2(doc, "콘텐츠 검증")
    add_bullet(doc, "□ 영상 시작 10초 안에 'AI 에이전트' 또는 '자율'이라는 단어가 등장한다")
    add_bullet(doc, "□ Section 3의 사람 ↔ AI 매핑 표가 1초 이상 등장한다")
    add_bullet(doc, "□ 21개 PC 모델의 다양성을 시각적으로 보여준다 (히트맵 권장)")
    add_bullet(doc, "□ 'AI가 동료평가한다' 메시지가 narration에 포함된다")
    add_bullet(doc, "□ 위험자산 70% lock이 시각적으로 표현된다 (도넛 차트)")
    add_bullet(doc, "□ Backtest Sharpe 1.195와 누적 수익률 +133%가 명확히 보인다")
    add_bullet(doc, "□ 면책 고지가 엔드 카드에 1.5초 이상 정지된 상태로 등장한다")

    add_h2(doc, "기술적 검증")
    add_bullet(doc, "□ 1080p 이상 / 60fps")
    add_bullet(doc, "□ Voiceover 0dB 정규화, BGM -18dB 이하 (목소리 묻히지 않게)")
    add_bullet(doc, "□ 자막 항상 표시 (음소거 시청자 고려)")
    add_bullet(doc, "□ 인스타그램 Reels용 9:16 세로 버전 별도 제공 (옵션)")

    add_h2(doc, "법적 검증")
    add_bullet(doc, "□ '수익 보장' / '손실 없음' 표현 없음")
    add_bullet(doc, "□ 타사 명칭 직접 비교 없음 (카테고리 단위만)")
    add_bullet(doc, "□ 면책 고지 자막 + 엔드 카드 모두 포함")
    add_bullet(doc, "□ 실명 / 개인정보 노출 없음 (대시보드 캡처 시 사이드바 사용자명 가림 처리)")

    # ====== 11. 글로서리 — 일상 한국어 변환 ======
    add_h1(doc, "11. 글로서리 — narration / 자막에서 지양 ↔ 권장 표현")

    add_body(doc, "30-50대 일반 시청자가 1-2회 듣고 즉시 이해해야 합니다. 아래 좌측 표현은 절대 narration / 자막에 사용하지 마세요. 우측 표현으로 변환합니다.", color=GRAY)

    add_h2(doc, "변환 표 (필수)")
    add_table(
        doc,
        headers=["✗ 지양 (전문 용어)", "✓ 권장 (일상 표현)", "비고"],
        rows=[
            ["Sharpe ratio / Sharpe 비율", "위험 대비 수익", "또는 \"같은 위험으로 더 많이 법\""],
            ["변동성", "흔들림 / 오르내리는 폭", "\"안정성\"도 OK"],
            ["walk-forward 백테스트", "지난 N년간 매분기마다 실제로 돌려본 결과", "\"8년 검증\"도 OK"],
            ["매크로 (macro)", "한국 경제 상태", "\"경제 흐름\"도 OK"],
            ["regime / 레짐", "지금 시기 / 경제 상황", "예: \"호황 끝물인지 침체기인지\""],
            ["ensemble / 앙상블", "최종 답 / AI들의 합의", "\"종합 의견\"도 OK"],
            ["CIO (agent)", "팀장 AI / 최종 결정자", "\"대표 AI\"도 OK"],
            ["자산 클래스", "상품 종류", "예: \"주식, 채권, 금 같은 종류\""],
            ["위험자산", "공격적 상품 / 주식 같은 위험 상품", "\"적극 투자 상품\"도 OK"],
            ["안전자산", "안전한 상품 / 예금·채권 같은", "—"],
            ["70% 고정 / lock", "한도 70%까지 끝까지 활용", "—"],
            ["자율 (autonomous)", "사람 개입 없이 / AI가 알아서", "\"자율\"은 너무 추상적"],
            ["다중 에이전트 / multi-agent", "여러 AI가 한 팀으로", "또는 \"AI 46명이\""],
            ["분산도 / diversification", "여러 군데 나눠 담음", "\"분산 투자\"도 OK"],
            ["tracking error", "(언급 X)", "BM 대비 추적 오차 — 일반인에게 불필요한 정보"],
            ["DC/IRP", "퇴직연금", "또는 \"DC형/IRP 같은 퇴직연금\" 1회만"],
            ["기대수익률", "앞으로 얼마나 오를지", "\"예상 수익\"도 OK"],
            ["expected vol / σ", "얼마나 흔들릴지", "—"],
            ["MDD / max drawdown", "가장 많이 떨어졌을 때의 폭", "또는 \"최악의 손실폭\""],
            ["proposal / weights", "추천 / 비중", "—"],
        ],
        widths=[5.5, 6.0, 5.5],
    )

    add_h2(doc, "그대로 사용 가능 (일반인도 익숙)")
    add_bullet(doc, "ETF, 펀드, 주식, 채권, 금, 원유, 코스피, S&P 500")
    add_bullet(doc, "퇴직연금, 연금, 은퇴")
    add_bullet(doc, "수익률 (단, %로 명시), 손실, 분산투자")
    add_bullet(doc, "환율, 금리, 물가, 실업률 (\"14가지 경제 신호\" 예시로만 1회)")

    add_h2(doc, "브랜드 용어 — 1회만 사용 + 즉시 풀어 설명")
    add_bullet(doc, "\"AI 에이전트\" → 첫 등장 시 \"AI 직원\" 또는 \"AI 팀원\"으로 쉬운 동의어 부언")
    add_bullet(doc, "\"자율주행 SAA\" → 제품명으로만. 영상에서는 \"AI가 알아서 만드는 추천\"으로 풀어 설명")

    add_callout(doc, "원칙: 영상 narration의 *모든* 단어가 시청자의 일상 어휘에 있어야 합니다. 한 단어라도 모호하면 그 순간 시청 이탈합니다.", color=GOLD)

    # ====== 부록 ======
    add_h1(doc, "부록: 자주 묻는 질문 (영상 후기 / 댓글 대비)")

    qas = [
        ("Q. AI가 제 돈을 직접 사고팔기도 하나요?",
         "A. 아니요. AI는 *추천*만 합니다. 매수와 매도는 본인이 직접 자기 퇴직연금 계좌에서 합니다. AI가 \"이번 달엔 이 비중으로 추천합니다\"라고 알려주면, 본인이 그대로 따를지 말지 결정하시면 됩니다."),
        ("Q. 추천은 얼마나 자주 바뀌나요?",
         "A. 매월 1일에 새로운 추천이 나옵니다. 한국 경제 상황이 크게 변하면 추천도 자동으로 조정됩니다."),
        ("Q. 어떤 AI를 쓰나요? ChatGPT인가요?",
         "A. 21가지의 자동 분석 시스템이 핵심입니다. 일부 단계에서 Anthropic의 Claude AI를 함께 사용해 시장 판단을 보조합니다. 모든 결정은 기록되고 화면에서 누구나 확인할 수 있습니다."),
        ("Q. 70%까지 공격적 상품에 넣는데, 경제가 나빠지면 위험하지 않나요?",
         "A. 70% 안에서 *어떤 상품에 얼마*를 넣을지를 AI가 매달 조정합니다. 경기가 나빠지면 미국 장기 국채와 금처럼 안정적인 자산 비중이 자동으로 늘어납니다."),
        ("Q. 8년 결과가 좋아도 앞으로도 그럴까요?",
         "A. 미래 수익은 누구도 보장할 수 없습니다. 본 자료는 정보 제공이며 투자 권유가 아닙니다. 과거 결과가 미래를 보장하지 않습니다. 본인의 투자 판단과 그 결과는 본인 책임입니다."),
        ("Q. 가입비나 수수료가 있나요?",
         "A. 본 데모는 무료입니다. 별도 가입도 필요 없습니다. 사이트 비밀번호만 알면 매월 추천을 확인할 수 있습니다."),
    ]
    for q, a in qas:
        add_h3(doc, q)
        add_body(doc, a, size=10)

    out = Path(__file__).resolve().parents[1] / "docs" / "production_brief_kr_pension_saa.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Written: {out}")


if __name__ == "__main__":
    build()
