"""
build_video_script.py
=====================
월간 연금 포트폴리오 유튜브 영상 스크립트 docx 생성.

대상: 40대 DC/IRP 연금 가입자
구성: 인트로(WHY) + 본론 5섹션 + 클로징 + 영상 설명란
분량: 약 9-11분

실행:
  cd d:\\파이선
  $env:PYTHONIOENCODING='utf-8'
  & "d:\\파이선\\pykrx_venv\\Scripts\\python.exe" -m \\
      kr_pension_portfolio.scripts.build_video_script \\
      --out outputs_trimmed10 --month 2026-05
"""
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
COLOR_ACCENT = RGBColor(0xC8, 0x9B, 0x3C)
COLOR_GREY = RGBColor(0x59, 0x59, 0x59)
COLOR_RED = RGBColor(0xC0, 0x39, 0x2B)
COLOR_GREEN = RGBColor(0x2E, 0x7D, 0x32)


def _set_run(run, *, bold=False, size=None, color=None, italic=False):
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if italic:
        run.italic = True


def add_title(doc, text, size=22, color=COLOR_PRIMARY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, bold=True, size=size, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, bold=True, size=15, color=COLOR_PRIMARY)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, bold=True, size=12, color=COLOR_GREY)
    return p


def add_para(doc, text, *, size=11, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, size=size, color=color, italic=italic, bold=bold)
    return p


def add_bullet(doc, text, *, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_run(r, size=size)
    return p


def add_quote_block(doc, text, *, size=12):
    """대본 본문(낭독 부분) — 들여쓰기 + 좌측 컬러바 느낌."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run(r, size=size, color=RGBColor(0x10, 0x10, 0x10))
    return p


def add_cue(doc, slide_no, text):
    """[화면 컷] 큐 — 슬라이드 번호 + 설명. 슬라이드 번호는 골드, 설명은 회색 italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run(f"🎬  [Slide {slide_no:02d}]  ")
    _set_run(r1, size=10, color=COLOR_ACCENT, bold=True)
    r2 = p.add_run(text)
    _set_run(r2, size=10, color=COLOR_GREY, italic=True)
    return p


def add_note(doc, text):
    """톤·전달 메모 — 보라색 italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(f"📝  {text}")
    _set_run(r, size=9, color=RGBColor(0x6A, 0x1B, 0x9A), italic=True)
    return p


def add_timestamp(doc, ts_label):
    p = doc.add_paragraph()
    r = p.add_run(f"[{ts_label}]")
    _set_run(r, bold=True, size=11, color=COLOR_ACCENT)
    return p


# ============================================================
# 메인
# ============================================================
def main(out_dir: str, month_label: str) -> None:
    base = Path(__file__).resolve().parent.parent
    out_path = base / out_dir

    cio = json.loads((out_path / "cio" / "final_portfolio.json").read_text(encoding="utf-8"))
    macro = json.loads((out_path / "macro" / "macro-view.json").read_text(encoding="utf-8"))
    cio_m = cio["metrics"]
    cio_w = cio["weights"]
    rd = macro["readings"]

    risk_w = (
        cio_w.get("kr-large-cap", 0) + cio_w.get("kr-dividend", 0)
        + cio_w.get("us-large-cap", 0) + cio_w.get("us-tech", 0)
        + cio_w.get("gold", 0)
    )

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ─── 표지 ───
    add_title(doc, f"월간 연금 포트폴리오 — {month_label}호 영상 스크립트", size=22)
    add_para(doc, "대상: 40대 DC/IRP 연금 가입자  |  분량: 약 9–11분  |  "
                  f"산출 기준일: {datetime.now().strftime('%Y-%m-%d')}",
             size=10, color=COLOR_GREY)
    add_para(doc, "구성: 인트로(WHY) → 매크로 진단 → CIO 추천 → 21모델 비교/강건성 "
                  "→ 위험 점검 → 다음 점검 → 클로징 + 설명란",
             size=10, color=COLOR_GREY)

    add_h2(doc, "핵심 메시지")
    add_bullet(doc, "21개 자산배분 모델이 합의한 자리를 매월 공개합니다 — 단일 모델 베팅 X")
    add_bullet(doc, "DC/IRP 70% 위험자산 한도는 알고리즘이 강제합니다 — 규제 준수 자동")
    add_bullet(doc, "그 달의 매크로 readings 14개를 함께 공개합니다 — 추적 가능성")
    add_bullet(doc, "정보 제공 자료, 매수 추천 아닙니다 — 면책 명시")

    add_h2(doc, "사용 가이드")
    add_para(doc, "• 본문 회색 박스(들여쓰기) = 낭독 대본",
             size=10, color=COLOR_GREY)
    add_para(doc, "• 🎬 [Slide NN] = 함께 표시할 슬라이드 번호 + 컷/인서트 설명",
             size=10, color=COLOR_GREY)
    add_para(doc, "• 📝 = 톤·전달 메모 (낭독 X)",
             size=10, color=COLOR_GREY)
    add_para(doc, f"• 슬라이드 데크: slides_{month_label}.pptx (총 29매, 16:9 widescreen)",
             size=10, color=COLOR_GREY)

    doc.add_page_break()

    # ============================================================
    # 인트로 — WHY
    # ============================================================
    add_h1(doc, "인트로 — 왜 월간 연금 포트폴리오인가")
    add_para(doc, "예상 분량: 약 1분 50초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_cue(doc, 1, "표지 (Slide 1) — 영상 시작 정지샷, 낭독 없음")
    add_cue(doc, 2, "챕터 1 디바이더 (Slide 2) — \"왜 월간 연금 포트폴리오인가\"")

    add_timestamp(doc, "0:00–0:15  훅")
    add_quote_block(doc,
        "40대 직장인이라면 한 번쯤 통장을 열어보고 이런 생각 해보셨을 겁니다.\n"
        "\"DC, IRP에 매달 돈은 들어가는데… 이걸 그냥 디폴트 옵션에 두는 게 정말 최선일까?\"\n"
        "그래서 직접 ETF를 골라보려고 하면 — 막막합니다.")
    add_cue(doc, 3, "통장 잔액 클로즈업 → 막막한 표정 → 물음표 그래픽")

    add_timestamp(doc, "0:15–0:55  문제 정의")
    add_quote_block(doc,
        "40대 연금 가입자가 마주치는 벽은 보통 세 가지입니다.\n\n"
        "첫째, 시간이 없습니다. 매일 시장을 들여다보고 분기마다 비중을 다시 잡을 여유가 없죠.\n\n"
        "둘째, 방법이 너무 많습니다. Equal Weight, Risk Parity, Max Sharpe, GMV — 들어보긴 했는데 "
        "어떤 게 정답인지 알 수가 없어요.\n\n"
        "셋째, 규제가 까다롭습니다. DC/IRP는 위험자산을 70%까지만 담을 수 있고, "
        "레버리지·인버스도 안 됩니다. 좋은 ETF를 발견해도 한도 안에 어떻게 끼워 넣을지 "
        "매번 계산해야 하죠.")
    add_cue(doc, 4, "3가지 벽 카드 — 시간 / 방법 / 규제")
    add_note(doc, "이 부분은 박자 늦춰 천천히 — 구독자가 자기 상황에 대입할 시간 필요.")

    add_timestamp(doc, "0:55–1:35  솔루션 컨셉")
    add_quote_block(doc,
        "그래서 이 채널은 매달 한 번, 자동화된 연금 포트폴리오 의견을 공개하기로 했습니다.\n\n"
        "핵심은 세 가지입니다.\n\n"
        "하나, 21개 서로 다른 자산배분 모델을 동시에 돌립니다. 어느 한 모델에 베팅하지 않고, "
        "21개가 어디서 합의하는지를 봅니다.\n\n"
        "둘, DC/IRP 70% 한도를 알고리즘이 강제합니다. 어떤 모델이 짠 비중이라도 한도를 넘지 않습니다.\n\n"
        "셋, 그 달의 매크로를 readings 14개로 함께 공개합니다. BOK 기준금리, KTB 10년물, "
        "USD/KRW, 브렌트유, VIX — 왜 이 비중이 나왔는지 추적할 수 있습니다.")
    add_cue(doc, 5, "솔루션 3가지 카드 (21모델 / 70% 한도 / 매크로 14개)")
    add_note(doc, "\"21개 모델이 합의하는 자리\" — 핵심 차별화. 굵게 강조 + 자막 키워드 노출.")

    add_timestamp(doc, "1:35–1:50  왜 월간인가")
    add_quote_block(doc,
        "분기마다 정기 리밸런싱이지만, 한 달 안에도 BOK 금통위, FOMC, 환율 같은 변수가 움직입니다. "
        "그래서 매달 점검합니다 — 분기까지 기다리지 말고, 매월 이 정도 비중이 합리적이다라는 "
        "기준선을 같이 보자는 취지입니다.")
    add_cue(doc, 6, "캘린더 — 5/6/7월 매월 점검, 8월 분기 리밸런싱")

    add_timestamp(doc, "1:50–2:00  전환")
    add_quote_block(doc,
        f"자, 그럼 {month_label.split('-')[1]}월 기준으로 21개 모델은 어떤 합의를 만들었는지, "
        "거시 진단부터 차근차근 보겠습니다.")
    add_cue(doc, 7, "챕터 2 디바이더 — \"5월 거시 진단\"")

    doc.add_page_break()

    # ============================================================
    # 본론 1 — 매크로 진단
    # ============================================================
    add_h1(doc, "본론 1 — 거시 진단: 5월 시장은 지금 어디에 있나")
    add_para(doc, "예상 분량: 약 1분 30초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_timestamp(doc, "2:00–2:15  도입")
    add_quote_block(doc,
        "어떤 비중을 짤지 결정하기 전에, 시장이 지금 어디에 와 있는지부터 봐야겠죠. "
        "저희 시스템은 매월 14개 매크로 지표를 ECOS와 FRED에서 자동으로 받아서 "
        "성장, 물가, 통화, 금융 — 이 네 축으로 점수화합니다.")
    add_cue(doc, 8, "데이터 소스 (ECOS · FRED) + 4축 점수 분석 도식")

    add_timestamp(doc, "2:15–2:45  성장 — 견조")
    add_quote_block(doc,
        f"성장부터 보면 — 지표가 굉장히 좋습니다.\n"
        f"GDP는 전년 대비 +{rd['kr_gdp_yoy']:.1f}%. 수출은 무려 +{rd['kr_exports_yoy']:.0f}% — "
        f"숫자만 보면 호황입니다. 실업률도 {rd['kr_unemployment']:.1f}%로 사실상 완전고용 수준이고요.\n"
        "다만 산업생산은 +1.2%로 외견상 호황 대비 미온적이라, "
        "수출이 일부 품목 집중일 가능성도 함께 봐야 합니다.")
    add_cue(doc, 9, "성장 4 KPI 카드 (GDP / 수출 / 실업 / 산업생산) + 초록 강조")

    add_timestamp(doc, "2:45–3:15  물가 — 안정, 그러나 유가 경계")
    add_quote_block(doc,
        f"물가는 헤드라인 CPI가 +{rd['kr_cpi_yoy']:.1f}%, 근원 CPI는 +{rd['kr_core_cpi_yoy']:.1f}% — "
        "한국은행 목표인 2% 근방에서 안정적으로 움직이고 있습니다.\n"
        f"하지만 안심은 이르고요. 브렌트유가 배럴당 ${rd['kr_brent_oil']:.0f}로 "
        "구조적 부담이 누적되고 있습니다. 유가가 여기서 한 단계 더 올라가면 "
        "성장은 멀쩡한데 물가만 다시 뜨는 — 스태그플레이션 위험이 살아납니다.")
    add_cue(doc, 10, "물가 KPI (CPI / 근원 CPI / Brent $114) + stagflation 경고 박스")
    add_note(doc, "\"스태그플레이션\"은 시청자가 잠깐 멈칫할 단어 — 0.5초 정도 여유.")

    add_timestamp(doc, "3:15–3:45  통화·금융 — BOK 인하 사이클")
    curve = macro.get("curve_signal", {})
    curve_line = (
        f"커브 상태는 {curve.get('regime')}이고, "
        f"{curve.get('curve_change_20d_bp', 0):+.0f}bp 변화가 관측됐습니다.\n"
        if curve else
        ""
    )
    add_quote_block(doc,
        f"통화 쪽은 BOK 기준금리가 {rd['kr_base_rate']:.2f}%까지 내려와 있습니다. "
        "작년부터 이어지는 인하 사이클이고요. "
        f"국고채 10년물은 {rd['kr_ktb_10y']:.2f}%, 3년물은 {rd['kr_ktb_3y']:.2f}% — "
        f"현재 3-10년 스프레드는 +{rd['kr_curve_3y_10y']*100:.0f}bp입니다.\n"
        f"{curve_line}"
        f"환율은 USD/KRW {rd['kr_usd_krw']:.0f}원, AA 회사채 스프레드 {rd['kr_corp_aa_spread_bp']:.0f}bp, "
        f"VIX {rd['us_vix']:.1f} — 금융 스트레스 지표는 모두 평온합니다.")
    add_cue(doc, 11, "통화·금융 KPI 8개 (BOK / KTB10Y / USD-KRW / VIX / 보조 4)")

    add_timestamp(doc, "3:45–4:00  결론")
    add_quote_block(doc,
        f"종합하면 — 시스템은 현재 시장을 '{macro['regime']}', "
        f"즉 경기 후반부로 분류했습니다. 신뢰도 {macro['confidence']:.2f}, "
        f"12개월 침체 확률은 {macro['recession_probability_12m']*100:.0f}% 수준입니다. "
        "성장은 좋은데 인플레와 유가 부담이 천천히 누적되는, "
        "전형적인 사이클 후반의 모습입니다.")
    add_cue(doc, 12, "4축 점수 막대 그래프 + late-cycle 레짐 카드")

    doc.add_page_break()

    # ============================================================
    # 본론 2 — CIO 추천
    # ============================================================
    add_h1(doc, "본론 2 — 5월 추천 포트폴리오")
    add_para(doc, "예상 분량: 약 1분 40초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_cue(doc, 13, "챕터 3 디바이더 — \"5월 추천 포트폴리오\"")

    add_timestamp(doc, "4:00–4:15  도입")
    add_quote_block(doc,
        f"이런 매크로 환경에서 21개 모델이 합의한 결과를 먼저 보여드리겠습니다. "
        f"채택된 ensemble은 inverse_te — 추적오차의 역수로 가중치를 준 합의 방식입니다.")
    add_cue(doc, 14, "21모델 → inverse_te ensemble → 최종 비중 흐름도")

    add_timestamp(doc, "4:15–5:00  Top 비중 5개")
    add_quote_block(doc,
        "비중을 큰 순서로 읽어드리면 —\n\n"
        f"첫째, KODEX KOFR금리액티브 {cio_w['kofr-cash']*100:.1f}%. 단기 자금시장 ETF로, "
        "BOK 기준금리에 거의 그대로 연동되는 현금 자산입니다.\n\n"
        f"둘째, KODEX 국고채10년 {cio_w['kr-treasuries-10y']*100:.1f}%. "
        "BOK 추가 인하 시 평가이익이 가장 크게 잡힐 듀레이션 자산입니다.\n\n"
        f"셋째, TIGER 머니마켓액티브 {cio_w['money-market']*100:.1f}%. "
        "또 다른 현금 분산용 ETF고요.\n\n"
        f"넷째, TIGER 미국채10년선물 {cio_w['us-treasuries-10y']*100:.1f}% — "
        "미국 듀레이션과 환노출을 동시에.\n\n"
        f"다섯째, TIGER 미국S&P500 {cio_w['us-large-cap']*100:.1f}% — "
        "글로벌 위험자산 코어 익스포저입니다.")
    add_cue(doc, 15, "10 ETF 비중 가로 막대 (내림차순) — 카테고리 색상")
    add_note(doc, "비중 숫자는 한 박자 멈추고 또박또박. 자막에 큰 글씨로 동시 노출.")

    add_timestamp(doc, "5:00–5:30  카테고리 합계")
    add_quote_block(doc,
        f"카테고리로 묶어보면 — 주식 {(cio_w['kr-large-cap']+cio_w['kr-dividend']+cio_w['us-large-cap']+cio_w['us-tech'])*100:.0f}%, "
        f"채권 {(cio_w['kr-treasuries-10y']+cio_w['us-treasuries-10y']+cio_w['us-ig-credit'])*100:.0f}%, "
        f"실물자산 {cio_w['gold']*100:.1f}%, 현금 {(cio_w['kofr-cash']+cio_w['money-market'])*100:.0f}% 입니다.\n"
        f"위험자산(주식 + 실물자산) 합계는 {risk_w*100:.1f}% — DC/IRP 한도 70%에서 "
        f"{(0.70-risk_w)*100:.0f}%포인트 여유가 있습니다. 시스템이 자율적으로 보수적인 자리에 머물러 있다는 뜻이고요.")
    add_cue(doc, 16, "카테고리 도넛 + 위험자산 38% 게이지 + 자산군 박스")

    add_timestamp(doc, "5:30–5:45  메트릭")
    add_quote_block(doc,
        f"이 비중의 기대수익률은 연 {cio_m['expected_return']*100:.2f}%, "
        f"기대변동성은 {cio_m['expected_vol']*100:.2f}%입니다. "
        f"백테스트 샤프지수는 {cio_m['backtest_sharpe']:.2f}, "
        f"최대 손실 폭은 {cio_m['backtest_maxdd']*100:.1f}% — "
        "한국 60/40 벤치마크가 같은 기간 -13.5% 빠진 것과 비교하면 절반 이하 손실입니다.")
    add_cue(doc, 17, "메트릭 4 카드 (E[r] / σ / Sharpe / MDD) + TE/HHI/EffN")

    doc.add_page_break()

    # ============================================================
    # 본론 3 — 21모델 비교 + 강건성
    # ============================================================
    add_h1(doc, "본론 3 — 21개 모델은 다 같은 결론을 내릴까?")
    add_para(doc, "예상 분량: 약 1분 45초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_cue(doc, 18, "챕터 4 디바이더 — \"21개 모델은 다 같은 결론?\"")

    add_timestamp(doc, "5:45–6:00  훅")
    add_quote_block(doc,
        "21개 모델은 다 같은 결론을 내릴까요? 답은 — 아닙니다. "
        "그런데 묘한 점이 있습니다.")
    add_cue(doc, 19, "타이포 슬라이드 — \"답은 아닙니다\" 강조 + \"묘한 점\"")

    add_timestamp(doc, "6:00–6:45  분포: 가장 공격적 vs 가장 방어적")
    add_quote_block(doc,
        "위험자산 비중만 봐도 21개 모델의 결론은 꽤 다릅니다.\n\n"
        "가장 공격적인 건 Market-Cap Weight 모델 — 위험자산 54.8%까지 채웠습니다. "
        "Equal Weight, Black-Litterman, CVaR-min, Max-Entropy 같은 모델들도 50% 라인까지 갑니다. "
        "IPS 한도 55% 직전이죠.\n\n"
        "가장 방어적인 건 Max-Drawdown-Constrained — 위험자산을 IPS 하한인 20%까지 줄여버렸습니다. "
        "TPA 22%, GMV 23.6%, Max-Sharpe 24.7% 같은 모델들도 비슷하게 보수적입니다.\n\n"
        "21개 평균은 38% 정도. 즉 동일한 IPS 제약 하에서도 위험자산 비중이 "
        "20%부터 55%까지 35%포인트 분포가 나오는 겁니다.")
    add_cue(doc, 20, "21모델 위험자산 비중 가로 막대 (정렬, 평균선·70% 한도선) + 공격/방어 박스")

    add_timestamp(doc, "6:45–7:30  강건성: ETF별 평균과 CIO가 일치")
    add_quote_block(doc,
        "그런데 — 이게 핵심인데요. "
        "ETF 단위로 21개 모델의 평균을 내봤습니다. 그리고 ensemble이 채택한 CIO 최종 비중과 비교했습니다.\n\n"
        "결과는 모든 ETF에서 1%포인트 이내로 일치했습니다.\n\n"
        "이게 무슨 뜻이냐면 — 어떤 PC 방법론을 쓰더라도, 5월 매크로와 CMA 환경에서는 "
        "결국 비슷한 자리로 수렴한다는 겁니다. 모델 의존성이 낮다는 뜻이고요. "
        "이 결과는 강건합니다 — 한두 모델의 가정이 틀렸다고 해서 결론이 뒤집히지 않습니다.")
    add_cue(doc, 21, "ETF별 21모델 평균 vs CIO 그룹 막대 + \"Δ ≤ 1pp\" 메시지 박스")
    add_note(doc, "\"강건합니다\"가 핵심 단어 — 톤 살짝 단단하게.")

    add_timestamp(doc, "7:30–7:45  마무리")
    add_quote_block(doc,
        "이게 매달 21개 모델을 동시에 돌리는 이유입니다. "
        "모델 합의가 강한 달은 \"이 자리는 단단하다\" — "
        "모델 합의가 깨지는 달은 \"무언가 변곡이 오고 있다\" — "
        "그 자체가 시그널이 됩니다.")
    add_cue(doc, 22, "합의/변곡 두 시나리오 — 좌(녹색) 합의 / 우(골드) 변곡")

    doc.add_page_break()

    # ============================================================
    # 본론 4 — 위험 점검
    # ============================================================
    add_h1(doc, "본론 4 — 위험 점검과 IPS 컴플라이언스")
    add_para(doc, "예상 분량: 약 1분 20초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_cue(doc, 23, "챕터 5 디바이더 — \"위험 점검\"")

    add_timestamp(doc, "7:45–8:00  통과 항목")
    add_quote_block(doc,
        "체크리스트 먼저 보겠습니다.\n"
        f"DC/IRP 위험자산 70% 한도 — 38%로 통과. "
        "카테고리 bounds도 모두 충족. 종목별 25% 한도도 최대 비중 14.3%라 여유 있게 통과합니다.")
    add_cue(doc, 24, "통과 항목 3 ✓ 박스 (한도/카테고리/종목)")

    add_timestamp(doc, "8:00–8:30  주의 항목")
    add_quote_block(doc,
        f"단, 두 가지 유의점이 있습니다.\n\n"
        f"첫째, 기대변동성이 {cio_m['expected_vol']*100:.2f}%로 IPS 하한 6%보다 낮습니다. "
        "위험이 너무 적다는 게 아니라 — 수익 기회를 충분히 잡지 못하는 자리라는 뜻입니다. "
        "사이클 후반부에 의도된 결과이긴 한데, 시청자께서 \"이거 너무 보수적인 거 아닌가\"라고 "
        "느끼셨다면 정확한 직관입니다.\n\n"
        f"둘째, 추적오차가 {cio_m['tracking_error']*100:.2f}%로 6% 예산을 넘었습니다. "
        "이건 한국 60/40 벤치마크가 KOSPI 60% + 국고채 40%로 너무 한국에 집중돼 있어서, "
        "글로벌 분산 포트폴리오와 자연스럽게 멀어진 결과입니다. "
        "다음 분기에 벤치마크 자체를 글로벌 60/40으로 재정의할 예정입니다.")
    add_cue(doc, 25, "주의 항목 2 ⚠ 박스 (σ 5.13% < 6% / TE 8.41% > 6%)")

    add_timestamp(doc, "8:30–9:00  스트레스 시나리오 3가지")
    add_quote_block(doc,
        "그리고 향후 한 달 안에 일어날 수 있는 스트레스 시나리오 세 가지를 미리 점검해 두겠습니다.\n\n"
        "하나, 유가 충격 — 브렌트가 $130을 넘으면, 골드 5%와 미국 IG 회사채 8%가 인플레 헷지로 작동합니다. "
        "현금 28%는 추가 매수 여력이고요.\n\n"
        "둘, BOK 추가 인하 — 2.5%에서 2.0%로 내려가면 "
        "한국 국고채 14%와 미국채 12%에서 평가이익 1.5~2% 정도 기대할 수 있습니다.\n\n"
        "셋, 환율 급등 — USD/KRW 1550 돌파 시 미국 주식 16% 환노출분에서 "
        "원화환산 수익률이 5~7% 추가됩니다.")
    add_cue(doc, 26, "스트레스 시나리오 3 카드 (유가 $130 / BOK 인하 / USD-KRW 1550)")

    doc.add_page_break()

    # ============================================================
    # 본론 5 — 다음 점검 포인트
    # ============================================================
    add_h1(doc, "본론 5 — 다음 한 달, 무엇을 보아야 하나")
    add_para(doc, "예상 분량: 약 50초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_timestamp(doc, "9:00–9:15  도입")
    add_quote_block(doc,
        "분기 정기 리밸런싱은 8월입니다. 그 사이 한 달 동안 시청자께서 "
        "함께 보시면 좋을 트리거 4가지를 정리해 드립니다.")
    add_cue(doc, 27, "다음 한 달 4 트리거 카드 (BOK / 유가 $120 / USD-KRW 1500·1450 / VIX 25)")

    add_timestamp(doc, "9:15–9:50  4가지 트리거")
    add_quote_block(doc,
        "하나, BOK 5월 금융통화위원회 결과. 추가 인하가 나오면 듀레이션 자산 추가 점검이 필요합니다.\n\n"
        "둘, 브렌트 유가가 $120을 돌파할지. 돌파 시 매크로 레짐이 stagflation으로 전환될 수 있습니다.\n\n"
        "셋, USD/KRW가 1500 또는 1450 어느 쪽으로든 벗어나는지. IPS 에스컬레이션 트리거입니다.\n\n"
        "넷, VIX가 25를 넘는지. 위험자산 추가 축소를 검토하는 자리입니다.")
    add_cue(doc, 27, "(같은 슬라이드) 4 트리거의 임계값 강조 — 카메라 줌인 효과 권장")

    # ============================================================
    # 클로징
    # ============================================================
    doc.add_paragraph()
    add_h1(doc, "클로징")
    add_para(doc, "예상 분량: 약 30초", size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    add_timestamp(doc, "9:50–10:10  마무리 + 면책")
    add_quote_block(doc,
        "오늘 5월호는 여기까지입니다.\n\n"
        "다시 한 번 강조드리면 — 본 영상은 자율주행 SAA 시스템이 산출한 결과를 "
        "정보로 공유드리는 자료입니다. 매수·매도 권유가 아니며, 실제 투자 결정 전 "
        "본인의 투자목적, 위험감수도, 재무상황을 종합적으로 검토하셔야 합니다.")
    add_cue(doc, 28, "면책 고지 슬라이드 (전문) — 잠시 정지 후 페이드")

    add_timestamp(doc, "10:10–10:30  CTA")
    add_quote_block(doc,
        "다음 6월호에서는 — BOK 5월 결정과 매크로 변화가 21개 모델 합의를 어떻게 "
        "흔들었는지 비교해 드리겠습니다.\n\n"
        "구독과 알림 설정 부탁드리고요. 댓글로 \"내 DC/IRP 통장에 이런 ETF가 있는데 "
        "어떻게 봐야 하나\" 같은 질문 주시면 다음 영상에서 다뤄보겠습니다.")
    add_cue(doc, 29, "다음 6월호 예고 + 구독 CTA — 따뜻한 톤")
    add_note(doc, "마지막 한 줄은 따뜻하게. 시청자가 \"내 통장 사정\"으로 자연스레 연결되도록.")

    doc.add_page_break()

    # ============================================================
    # 영상 설명란 (Description)
    # ============================================================
    add_h1(doc, "영상 설명란 (YouTube Description)")
    add_para(doc, "그대로 복사해서 영상 업로드 시 설명란에 붙이세요.",
             size=10, color=COLOR_GREY, italic=True)
    doc.add_paragraph()

    desc = (
        f"📊 월간 연금 포트폴리오 — {month_label}호\n"
        "DC/IRP 가입자를 위한 자동화 SAA 의견. 21개 자산배분 모델이 "
        "합의한 자리를 매월 공개합니다.\n\n"
        "─ 이번 달 핵심 ─\n"
        f"• 거시 레짐: {macro['regime']} (신뢰도 {macro['confidence']:.2f}, "
        f"12m 침체확률 {macro['recession_probability_12m']*100:.0f}%)\n"
        f"• 채택 ensemble: inverse_te\n"
        f"• 위험자산 비중: {risk_w*100:.1f}% (DC/IRP 한도 70%)\n"
        f"• 기대수익률 {cio_m['expected_return']*100:.2f}% / σ {cio_m['expected_vol']*100:.2f}%\n"
        f"• BT Sharpe {cio_m['backtest_sharpe']:.2f} / MDD {cio_m['backtest_maxdd']*100:.2f}%\n\n"
        "─ 챕터 ─\n"
        "0:00 인트로 — 왜 월간 연금 포트폴리오인가\n"
        "2:00 5월 거시 진단\n"
        "4:00 21개 모델이 합의한 추천\n"
        "5:45 21개 모델은 다 같은 결론?\n"
        "7:45 위험 점검\n"
        "9:00 다음 한 달 점검 포인트\n"
        "9:50 마무리\n\n"
        "─ 매크로 데이터 출처 ─\n"
        "• 한국: 한국은행 ECOS API (GDP, CPI, 실업, 수출, BOK 금리, KTB)\n"
        "• 미국: FRED (Fed funds, VIX, Brent oil)\n\n"
        "─ 면책 고지 ─\n"
        "본 영상은 자율주행 SAA 파이프라인이 산출한 정보를 시각화한 정보 제공 자료이며, "
        "특정 금융상품의 매수·매도 권유가 아닙니다. 투자자문업·투자일임업·투자권유 행위가 "
        "아닙니다. 백테스트 결과는 가상의 가정에 기반하며 미래 성과를 보장하지 않습니다. "
        "DC/IRP 규제(위험자산 70% 한도, 레버리지·인버스 ETF 매수 금지)를 가정한 IPS "
        "하에서 산출되었습니다. 실제 투자 결정 전 본인의 투자목적, 위험감수도, 재무상황을 "
        "종합 검토하시기 바랍니다.\n\n"
        "#연금 #DC #IRP #자산배분 #ETF #포트폴리오 #자율주행 #SAA"
    )
    p = doc.add_paragraph()
    r = p.add_run(desc)
    _set_run(r, size=10)

    # 저장
    save_path = out_path / f"video_script_{month_label}.docx"
    doc.save(save_path)
    print(f"[saved] {save_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="outputs_trimmed10")
    parser.add_argument("--month", default=None)
    args = parser.parse_args()
    month_label = args.month or datetime.now().strftime("%Y-%m")
    main(out_dir=args.out, month_label=month_label)
