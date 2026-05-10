"""Generate standalone narration script docx for voiceover / subtitle production.

Output: docs/narration_script_kr_pension_saa.docx
Audience: Korean voiceover artist / 자막 작성자
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xC8, 0x9B, 0x3C)
GRAY = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x22, 0x22, 0x22)


def set_kr_font(run, size_pt=11, bold=False, color=None, italic=False):
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    rpr.append(rfonts)


def shade_para(p, hex_color):
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_segment(doc, timecode, length_sec, narration, screen, voiceover_note=""):
    """Render one narration segment in a script-friendly layout."""
    # 시간/길이 헤더
    p = doc.add_paragraph()
    shade_para(p, "1F4E79")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"  {timecode}   ·   {length_sec}초  ")
    set_kr_font(r, 11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 화면 큐 (회색 배경)
    p = doc.add_paragraph()
    shade_para(p, "F2F2F2")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    r = p.add_run("화면: ")
    set_kr_font(r, 9.5, bold=True, color=GRAY)
    r = p.add_run(screen)
    set_kr_font(r, 9.5, color=GRAY)

    # Narration (대형 텍스트)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(narration)
    set_kr_font(r, 13, color=BLACK)

    # 성우 메모 (선택)
    if voiceover_note:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("성우 톤: ")
        set_kr_font(r, 9.5, bold=True, color=GOLD)
        r = p.add_run(voiceover_note)
        set_kr_font(r, 9.5, italic=True, color=GRAY)

    # 자막 길이 가이드
    char_count = len(narration.replace(" ", ""))
    cps = char_count / max(length_sec, 1)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"음절 수: {char_count}자  ·  발화 속도: {cps:.1f}자/초  ·  자막은 narration의 80% 길이로 단축")
    set_kr_font(r, 8.5, color=GRAY, italic=True)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_kr_font(r, 18, bold=True, color=NAVY)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_kr_font(r, 13, bold=True, color=color)


def add_body(doc, text, size=10, color=None, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_kr_font(r, size, bold=bold, color=color)


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_kr_font(r, size)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---------- 표지 ----------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("NARRATION SCRIPT")
    set_kr_font(r, 11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KR 연금 자율주행 SAA")
    set_kr_font(r, 28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("홍보 영상 — 성우/자막 작업용")
    set_kr_font(r, 13, color=GRAY)
    p.paragraph_format.space_after = Pt(40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("메인 3분 30초 + 티저 60초")
    set_kr_font(r, 14, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(40)

    # 작업 노트
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📋 성우 / 자막 작성 시 주의사항")
    set_kr_font(r, 11, bold=True, color=NAVY)

    note_lines = [
        "• 모든 narration은 30-50대 일반 시청자가 1-2회 듣고 즉시 이해하도록 일상 한국어로 작성됨.",
        "• 자막은 narration의 80% 길이로 단축 (시청자가 읽으면서 듣는 부담 감소).",
        "• \"흔들림\", \"위험 대비 수익\" 같은 일상 표현이 의도된 것 — 전문 용어로 \"고치지\" 말 것.",
        "• 시간 정합성을 위해 단어 가감은 가능하나 핵심 표현은 유지.",
        "• 모든 segment 끝에 명시된 음절 수 / 발화 속도 가이드 참고.",
    ]
    for line in note_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(2.5)
        p.paragraph_format.right_indent = Cm(2.5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_kr_font(r, 10, color=GRAY)

    doc.add_page_break()

    # ---------- 1. 메인 영상 (3:30) ----------
    add_h1(doc, "메인 영상 — 3분 30초 (사이드바 데모 흐름)")

    add_body(doc, "구조: Hook → AI 경제 분석가 → AI 자산 분석가 18명 → AI 매니저 21명 → AI 동료평가 + 팀장 결정 → 이번 달 추천 → 8년 검증 → CTA. 사이드바 메뉴 2 → 3 → 5 → 6 → 1 순으로 화면 이동.", size=10, color=GRAY)

    main_segs = [
        ("0:00 — 0:30", 30,
         "퇴직연금, 신경 쓰고 계신가요? 한국 퇴직연금은 주식처럼 공격적인 상품에 70%까지 넣을 수 있습니다. 그런데 대부분 그 한도를 다 쓰지 못합니다. 사실, 8년이면 거의 100% 가까운 차이가 납니다. 이걸 AI에게 맡기면 어떨까요? 46명의 AI가 한 팀이 되어, 매달 추천 포트폴리오를 만들어준다면.",
         "안전자산 100% 누적곡선 vs 60/40 누적곡선 차이를 강조하는 차트 (10초) → AI 46명 매핑 표 페이드 인 (15초) → 본 영상 타이틀 (5초)",
         "차분하게 시작 → 후반 \"이걸 AI에게 맡기면 어떨까요?\"에서 살짝 호기심 어린 톤"),

        ("0:30 — 0:50", 20,
         "먼저, AI가 지금 한국 경제 상태부터 살핍니다. 금리, 물가, 환율, 수출, 실업률... 14가지 경제 신호를 모아서, 지금이 좋은 시기인지 조심할 시기인지 판단합니다. 사람이 하면 며칠 걸릴 일을 AI는 자동으로 합니다.",
         "사이드바 메뉴 \"2_Macro_Regime\" 클릭 (1초) → regime 라벨 클로즈업 (예: \"late-cycle, 신뢰도 0.53\") → 14개 지표 표 빠른 컷 (GDP, CPI, BOK 금리, USD/KRW 등 강조)",
         "신뢰감 + 정보 제공 톤. \"14가지\" 같은 숫자는 살짝 강조"),

        ("0:50 — 1:10", 20,
         "다음, 18명의 AI가 각자 한 가지 ETF를 맡습니다. 코스피 ETF는 1번 AI가, 미국 S&P 500은 2번 AI가, 미국 국채는 3번 AI가... 각자 자기 담당 상품이 앞으로 얼마나 오를지, 얼마나 흔들릴지 예측합니다. 18명이 동시에, 24시간 일합니다.",
         "사이드바 메뉴 \"3_Asset_Classes\" 클릭 → 18개 ETF 카드 빠른 컷 (각 ~1초, ETF명 + E[r] 강조)",
         "리듬감 있게. \"코스피는, S&P는, 국채는...\" 부분에서 약간 빠른 페이스로 다양성 강조"),

        ("1:10 — 1:35", 25,
         "이제 21명의 AI 매니저가 등장합니다. 같은 자료를 받지만, 생각하는 방식이 다릅니다. 어떤 AI는 가장 안전한 답, 어떤 AI는 가장 수익이 클 답을 추천합니다. 21명이 21개의 서로 다른 추천을 냅니다. 이 화면이 그 차이를 보여줍니다 — 색이 진할수록 비중이 큰 상품.",
         "사이드바 메뉴 \"5_월간_연금_포트폴리오\" 클릭 → 21×ETF 히트맵 zoom-in. 색상 차이 시각화 (예: 한 모델은 미국주식 25%, 다른 모델은 채권 20% 같이 대비)",
         "\"같은 자료를 받지만, 생각하는 방식이 다릅니다\"가 hook 문장. 살짝 멈춰주고 강조"),

        ("1:35 — 1:55", 20,
         "AI들끼리 서로 점수를 매깁니다. 누구의 추천이 가장 믿을 만한가? 가장 높은 점수를 받은 5명을 추려내고, 팀장 AI가 지금 시장에 가장 잘 맞는 답을 최종 선택합니다. 사람의 개입 없이, 모든 평가가 자동으로 이루어집니다.",
         "같은 페이지 스크롤 다운 → Borda 투표 표 (Top-5 강조) → CIO 최종 도넛으로 트랜지션",
         "신뢰감 + 자율성 강조. 마지막 \"사람의 개입 없이\" 부분에서 약간 무게 있게"),

        ("1:55 — 2:15", 20,
         "이번 달 추천 포트폴리오입니다. 주식 55%에 금과 원유가 15%, 합쳐서 공격적 상품 70%. 나머지 30%는 채권과 현금. 법이 허용하는 한도까지 끝까지 활용하는 게 핵심입니다.",
         "사이드바 메뉴 \"6_현재_포트폴리오\" 클릭 → 도넛 차트 (55+15=70 강조 표시) + 18 ETF 비중 표 상위 5개 클로즈업",
         "단호하고 명료하게. 결과를 \"보여주는\" 톤"),

        ("2:15 — 3:00", 45,
         "근데 진짜 효과 있을까요? 지난 8년, 2018년부터 지금까지, 매분기마다 이 AI 시스템을 실제로 돌려봤습니다. 결과는 — 한국에서 가장 흔한 60/40 펀드와 비교했을 때, 흔들림은 30% 작고, 위험 대비 수익은 30% 더 좋았습니다. 같은 8년에 100만원을 넣었으면 233만원이 됩니다.",
         "사이드바 메뉴 \"1_Backtest\" 클릭 → 변형 선택에서 \"v1 lock70 Baseline\" 클릭 → NAV 곡선 (lock70 vs 60/40 BM 두 줄) 풀스크린 → 메트릭 카드 4장 (Sharpe / 수익률 / 흔들림 / 최대 낙폭) 빠른 컷",
         "\"근데 진짜 효과 있을까요?\"는 의심 어린 톤으로 → \"100만원이 233만원\" 부분은 결정적/단호하게"),

        ("3:00 — 3:30", 30,
         "매월 1일, 새로운 추천이 자동으로 나옵니다. 사이트에서 바로 확인할 수 있습니다. 무료입니다. 본 자료는 정보 제공이며, 투자 권유가 아닙니다.",
         "URL https://kr-pension-saa-gloual88.streamlit.app/ + QR 코드 + 비밀번호 3734 (5초) → 면책 고지 풀스크린 (5초 정지) → 로고/엔드 카드",
         "차분하고 신뢰감 있게 마무리. 면책 부분은 약간 더 천천히, 분명히"),
    ]

    for tc, sec, narr, screen, vo in main_segs:
        add_segment(doc, tc, sec, narr, screen, vo)

    doc.add_page_break()

    # ---------- 2. 60초 티저 ----------
    add_h1(doc, "티저 — 60초 (메인의 압축)")

    add_body(doc, "용도: YouTube Shorts / Instagram Reels / 카카오톡 채널. 빠른 컷 + 강한 후크.", size=10, color=GRAY)

    teaser_segs = [
        ("0:00 — 0:10", 10,
         "퇴직연금. 한도 70%, 안 쓰면 손해입니다.",
         "안전자산 vs 60/40 누적 차이 차트 (3초) → AI 46명 매핑 표 페이드 인 (5초) → 영상 타이틀 (2초)",
         "강하고 직접적. \"안 쓰면 손해\" 부분에서 약간 무게 있게"),

        ("0:10 — 0:25", 15,
         "그래서 만들었습니다. 46명의 AI 팀. 한국 경제부터 ETF 비중까지 — 자기들끼리 의논해서 매달 추천을 만듭니다.",
         "사이드바 메뉴 빠른 컷: Macro_Regime → Asset_Classes → 21 PC 히트맵 → CIO 도넛. 각 3-4초, 마우스 이동 부드럽게",
         "흥미진진하게. \"그래서 만들었습니다\"에서 짧은 멈춤"),

        ("0:25 — 0:45", 20,
         "지난 8년 검증 결과: 60/40 펀드 대비 흔들림 30% 작고, 수익도 더 좋았습니다.",
         "Backtest 페이지 NAV 곡선 (lock70 vs BM 두 줄) → Sharpe 1.195 메트릭 카드 → 누적 +133% 카드. 각 5초씩 빠른 컷",
         "결정적/확신 있게. \"30% 작고\"와 \"수익도 더 좋았습니다\" 두 번 강조"),

        ("0:45 — 1:00", 15,
         "사이트에서 무료. 매월 1일 자동. 본 자료는 투자 권유가 아닙니다.",
         "URL + QR + 비밀번호 3734 + 면책 고지 (정지)",
         "차분하게 마무리. 면책 분명히 발화"),
    ]

    for tc, sec, narr, screen, vo in teaser_segs:
        add_segment(doc, tc, sec, narr, screen, vo)

    doc.add_page_break()

    # ---------- 3. 발화 가이드라인 ----------
    add_h1(doc, "발화 가이드라인")

    add_h2(doc, "전반적 톤")
    add_bullet(doc, "신뢰감 + 데이터 기반의 차분함. \"권유\"가 아니라 \"설명\".")
    add_bullet(doc, "30-40대 남성 또는 여성 모두 가능. 전문성이 우선이며 친근감은 부차.")
    add_bullet(doc, "Hook 막은 약간 호기심을, Result 막은 결정적으로, CTA는 차분하게.")

    add_h2(doc, "발화 속도")
    add_bullet(doc, "기본: 분당 280-320 음절 (4.5-5.3 음절/초)")
    add_bullet(doc, "각 segment 끝의 \"음절 수 / 발화 속도\" 가이드 준수")
    add_bullet(doc, "숫자(70%, 30%, 100만원 → 233만원)는 살짝 천천히 강조")

    add_h2(doc, "발음 주의 단어")
    add_table_2col(doc, [
        ("ETF", "이티에프 (영어 그대로 읽기)"),
        ("S&P 500", "에스앤피 오백"),
        ("코스피", "코스피 (그대로)"),
        ("AI", "에이아이"),
        ("46명", "사십육 명"),
        ("100% 가까운", "백 퍼센트 가까운"),
        ("10년물", "십년물"),
        ("100만원이 233만원", "백만 원이 이백 삼십삼 만 원"),
    ])

    add_h2(doc, "강조 단어 (살짝 힘주기)")
    add_bullet(doc, "\"46명의 AI\" — 차별성의 핵심")
    add_bullet(doc, "\"각자 다른 방식\" — 21개 모델 다양성 강조")
    add_bullet(doc, "\"사람의 개입 없이\" — 자율성 강조")
    add_bullet(doc, "\"흔들림 30% 작고, 수익은 더 좋았습니다\" — Result 핵심 메시지")
    add_bullet(doc, "\"100만원이 233만원\" — 구체성으로 신뢰감")

    add_h2(doc, "회피해야 할 표현 (자막 작성 시 주의)")
    add_bullet(doc, "\"수익이 보장됩니다\" / \"손실이 없습니다\" — 절대 사용 금지 (법적 위험)")
    add_bullet(doc, "\"전문가가 운용합니다\" — 본 모델은 AI 자율, 사람 자문이 아님")
    add_bullet(doc, "\"지금 가입하세요\" — 무료 데모이므로 \"지금 확인하세요\"가 적절")

    doc.add_page_break()

    # ---------- 4. 면책 고지 ----------
    add_h1(doc, "면책 고지 (필수 자막 / 엔드 카드)")

    add_body(doc, "아래 문구는 영상 엔드 카드(마지막 5초)에 풀스크린으로 표시되어야 합니다. 자막은 평생 정주행 가능하도록 명확하게.", color=GRAY)

    p = doc.add_paragraph()
    shade_para(p, "F8F8F8")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.6
    r = p.add_run(
        "「본 자료는 정보 제공 목적이며 투자 권유가 아닙니다. "
        "투자 의사결정과 그 결과는 투자자 본인의 책임입니다. "
        "백테스트 성과는 미래 수익을 보장하지 않습니다. "
        "표시된 추천 포트폴리오는 예시이며, 본인 상황에 맞는지는 별도로 판단하시기 바랍니다.」"
    )
    set_kr_font(r, 11, color=BLACK)

    add_body(doc, "Voiceover 추가 권장 (3:00-3:30 CTA 막에 narration으로 포함):", size=10, color=GRAY)
    add_body(doc, "「본 자료는 정보 제공이며, 투자 권유가 아닙니다.」", size=11, bold=True)

    out = Path(__file__).resolve().parents[1] / "docs" / "narration_script_kr_pension_saa.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Written: {out}")


def add_table_2col(doc, rows):
    """Helper: 발음 표 (2열 표)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        set_kr_font(r0, 10, bold=True)
        r1 = c1.paragraphs[0].add_run(v)
        set_kr_font(r1, 10, color=GRAY)
    doc.add_paragraph()


if __name__ == "__main__":
    build()
